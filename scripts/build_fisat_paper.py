import re
from pathlib import Path
from shutil import copyfile
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "submission_fisat"
LATEX_TEMPLATE_DIR = Path("D:/Springer_Latex_Template")
DOCX_TEMPLATE = Path("D:/Springer_DOCX_Template/splnproc1703.docm")

TITLE = (
    "VaccineNLP: An Explainable Human-in-the-Loop Multi-Task NLP "
    "Framework for Vietnamese Vaccine Misinformation Surveillance"
)
TITLE_RUNNING = "VaccineNLP: Explainable HITL Misinformation Surveillance"
AUTHORS = [
    {"name": "Manh Hung Kim", "inst": "1"},
    {"name": "Quynh Phuong Dinh Le", "inst": "1"},
    {"name": "Lam Quan Tran", "inst": "2"},
    {"name": "Hong Viet Tran", "inst": "3"},
    {"name": "Hang Nguyet Van Nguyen", "inst": "1"},
]
AUTHOR_RUNNING = "Manh Hung Kim et al."
INSTITUTIONS = [
    "Ha Noi University of Public Health, Ha Noi, Viet Nam",
    "Digital Health Center, Ha Noi University of Public Health, Ha Noi, Viet Nam",
    (
        "Institute of Artificial Intelligence, VNU University of Engineering "
        "and Technology, Vietnam National University, Hanoi, Vietnam"
    ),
]
CONTACT_SECTION_TITLE = "Correspondence and Author Contacts"
CONTACT_EMAILS = [
    "211090016@studenthuph.edu.vn",
    "quantl3@fpt.edu.vn",
    "thv79@gmail.com",
]
CONTACT_PARAGRAPHS = [
    (
        "Correspondence concerning this paper should be addressed to the first author, "
        "Manh Hung Kim, Ha Noi University of Public Health, Ha Noi, Viet Nam. "
        "Email: 211090016@studenthuph.edu.vn."
    ),
    (
        "Additional author contact: Lam Quan Tran, Digital Health Center, Ha Noi "
        "University of Public Health, Ha Noi, Viet Nam. "
        "Email: quantl3@fpt.edu.vn."
    ),
    (
        "Additional author contact: Hong Viet Tran, Institute of Artificial Intelligence, "
        "VNU University of Engineering and Technology, Vietnam National University, "
        "Hanoi, Vietnam. Email: thv79@gmail.com."
    ),
]

ABSTRACT = (
    "Vaccine misinformation has become a public-health informatics challenge "
    "in Vietnamese digital environments, where social media posts, online "
    "comments, and informal vaccine narratives can shape vaccine confidence. "
    "This paper presents VaccineNLP, an explainable human-in-the-loop "
    "multi-task natural language processing framework for Vietnamese vaccine "
    "misinformation surveillance. We collected 1,856 Vietnamese vaccine-related "
    "texts from Facebook, YouTube, online news, Reddit, and public forums during "
    "2020--2026, then applied an eight-step Vietnamese preprocessing pipeline "
    "and a three-axis annotation schema covering misinformation, stance, and "
    "sentiment. A 186-sample expert-validated Gold Test Set was used to compare "
    "XLM-RoBERTa, PhoBERT-v2 multi-task learning, and a Gemma-4 4B QLoRA "
    "reasoning model. PhoBERT-v2 achieved the highest average Macro F1 (0.6967), "
    "while XLM-RoBERTa obtained the best misinformation Macro F1 (0.7038) and "
    "Gemma-4 4B obtained the best sentiment Macro F1 (0.7700). Temperature "
    "Scaling reduced Expected Calibration Error by 44--56% across the three "
    "axes. However, agreement between the LLM annotator and expert labels was "
    "very low for misinformation (Cohen's kappa = 0.0525), indicating that LLMs "
    "should support, not replace, expert medical fact-checking. Statistical "
    "tests further showed that anti-vaccine stance and negative sentiment were "
    "strongly associated with misinformation risk. The results support a cautious "
    "trustworthy-AI design for Vietnamese infodemic monitoring."
)

KEYWORDS = [
    "Vaccine misinformation",
    "Vietnamese NLP",
    "Infodemic surveillance",
    "Human-in-the-loop",
    "Explainable AI",
    "Multi-task learning",
    "PhoBERT",
    "Confidence calibration",
]

FIGURES = {
    "architecture": {
        "source_doc": Path("D:/Bao_cao_toan_van_VaccineNLP.docx"),
        "source_media": "word/media/image1.png",
        "filename": "fig1_architecture_en.png",
        "caption": "VaccineNLP two-tier human-in-the-loop architecture for Vietnamese vaccine misinformation surveillance.",
        "label": "fig:architecture",
        "latex_width": r"0.78\textwidth",
        "docx_width_cm": 12.1,
        "alt": "VaccineNLP architecture diagram showing preprocessing, PhoBERT-v2 classification, Gemma-4 explanation, RAG evidence checking, and cognitive integrity output.",
        "generated": "architecture",
    },
    "per_class_f1": {
        "source_doc": Path("D:/Bao_cao_toan_van_VaccineNLP.docx"),
        "source_media": "word/media/image2.png",
        "filename": "fig2_per_class_f1_en.png",
        "caption": "Per-class F1 scores of XLM-RoBERTa, PhoBERT-v2, and Gemma-4 4B on the Gold Test Set.",
        "label": "fig:perclass",
        "latex_width": r"\textwidth",
        "docx_width_cm": 12.2,
        "alt": "Bar charts comparing per-class F1 for misinformation, stance, and sentiment across XLM-RoBERTa, PhoBERT-v2, and Gemma-4 4B.",
        "generated": "per_class_f1",
    },
    "calibration": {
        "source_doc": Path("D:/Bao_cao_toan_van_VaccineNLP.docx"),
        "source_media": "word/media/image3.png",
        "filename": "fig3_calibration_en.png",
        "caption": "PhoBERT-v2 reliability diagrams before and after Temperature Scaling across the three tasks.",
        "label": "fig:calibration",
        "latex_width": r"\textwidth",
        "docx_width_cm": 12.2,
        "alt": "Reliability diagrams showing raw and calibrated confidence versus accuracy for misinformation, stance, and sentiment.",
        "generated": "calibration",
    },
    "prototype": {
        "source_doc": Path("D:/DATN_2211090016_2211090031_BanClean_DaSuaBia_SoTrangTren (1).docx"),
        "source_media": "word/media/image8.jpeg",
        "filename": "fig4_prototype_stream.jpeg",
        "caption": "VaccineNLP prototype interface for conversation-flow screening and evidence-oriented review.",
        "label": "fig:prototype",
        "latex_width": r"\textwidth",
        "docx_width_cm": 12.2,
        "alt": "Screenshot of the VaccineNLP web prototype showing conversation-flow analysis, risk indicators, contextual evidence, and review output.",
    },
}

TABLES = {
    "sources": {
        "caption": "Data sources and cleaning pipeline.",
        "label": "tab:sources",
        "headers": [
            "Source",
            "Initial",
            "Deduplicated",
            "Cleaned",
            "Annotated",
            "%",
        ],
        "rows": [
            ["Online news", "436", "338", "291", "270", "14.5"],
            ["Facebook", "1,212", "922", "812", "750", "40.4"],
            ["YouTube", "1,168", "890", "784", "723", "39.0"],
            ["Reddit/forums", "184", "140", "123", "113", "6.1"],
            ["Total", "3,000", "2,290", "2,010", "1,856", "100.0"],
        ],
        "align": ["left", "center", "center", "center", "center", "center"],
        "latex_align": "lrrrrr",
        "widths": [1700, 900, 1150, 900, 1000, 650],
        "note": (
            "Annotated samples are texts retained after source-level filtering, "
            "deduplication, text cleaning, and vaccine-relevance screening."
        ),
    },
    "literature_comparison": {
        "caption": "Research-positioning comparison of VaccineNLP against related misinformation and vaccine NLP studies.",
        "label": "tab:literature",
        "headers": ["Study", "Setting", "Main task / model", "Reported result", "Gap addressed by VaccineNLP"],
        "rows": [
            [
                "ANTi-Vax [5]",
                "English Twitter vaccine comments",
                "Binary vaccine misinformation detection with transformer classifiers",
                "High in-domain F1 on a large English corpus",
                "Not Vietnamese; binary labels do not separate misinformation, stance, and sentiment.",
            ],
            [
                "VaxxHesitancy [25]",
                "English Twitter vaccine hesitancy",
                "Domain-specific VaxxBERT for stance and hesitancy analysis",
                "F1 about 0.69 for hesitancy classification",
                "Focuses on hesitancy rather than medical misinformation verification.",
            ],
            [
                "Checkovid [26]",
                "COVID-19 misinformation on Twitter",
                "Network and content mining with ML/NLP classifiers",
                "Strong F1 under Twitter-based COVID-19 detection settings",
                "Relies on Twitter-centric features; not tailored to Vietnamese vaccine discourse.",
            ],
            [
                "DANN + LIME on CoAID/MiSoVac [27]",
                "Cross-platform COVID-19 misinformation",
                "Domain adaptation plus explainable AI",
                "Improved cross-domain F1 with LIME explanations",
                "Explanation is model-agnostic; no Vietnamese HITL Gold Test Set or calibration analysis.",
            ],
            [
                "LLM annotation studies [13]",
                "General text annotation tasks",
                "LLMs as labelers or crowd-worker alternatives",
                "Often competitive in non-medical annotation",
                "Medical truth claims still require expert review; our misinformation kappa was only 0.0525.",
            ],
            [
                "VaccineNLP (this study)",
                "Vietnamese multi-source vaccine discourse",
                "XLM-R, PhoBERT-v2 multi-task learning, Gemma-4 4B QLoRA, HITL, calibration",
                "Mean Macro F1 = 0.6967; ECE reduced 44--56%",
                "Provides Vietnamese tri-axis labels, expert validation, calibrated confidence, and operational XAI.",
            ],
        ],
        "align": ["left", "left", "left", "left", "left"],
        "latex_align": "lllll",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.16\textwidth}"
            r">{\raggedright\arraybackslash}p{0.17\textwidth}"
            r">{\raggedright\arraybackslash}p{0.22\textwidth}"
            r">{\raggedright\arraybackslash}p{0.17\textwidth}"
            r">{\raggedright\arraybackslash}p{0.20\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "1.8pt",
        "arraystretch": "1.08",
        "widths": [1050, 1200, 1650, 1350, 1350],
        "font_size": 6.8,
        "note": "Reported scores are not directly comparable because datasets, languages, label definitions, and validation protocols differ.",
    },
    "labels": {
        "caption": "Gold Test Set label distribution (n = 186).",
        "label": "tab:labels",
        "headers": ["Axis", "Class", "n", "%"],
        "rows": [
            ["Misinformation", "Correct / not misleading", "158", "84.9"],
            ["Misinformation", "Misinformation", "28", "15.1"],
            ["Stance", "Supportive", "54", "29.0"],
            ["Stance", "Opposing / hesitant", "48", "25.8"],
            ["Stance", "Neutral", "84", "45.2"],
            ["Sentiment", "Negative", "71", "38.2"],
            ["Sentiment", "Neutral", "75", "40.3"],
            ["Sentiment", "Positive", "40", "21.5"],
        ],
        "align": ["left", "left", "center", "center"],
        "latex_align": "llrr",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.28\textwidth}"
            r">{\raggedright\arraybackslash}p{0.46\textwidth}"
            r">{\centering\arraybackslash}p{0.08\textwidth}"
            r">{\centering\arraybackslash}p{0.08\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.5pt",
        "arraystretch": "1.06",
        "widths": [1600, 3200, 900, 900],
    },
    "schema": {
        "caption": "Operational annotation schema used for HITL review.",
        "label": "tab:schema",
        "headers": ["Axis", "Label", "Operational definition"],
        "rows": [
            [
                "Misinformation",
                "Misinformation",
                "Contains false, misleading, unsupported, or context-stripped vaccine claims.",
            ],
            [
                "Misinformation",
                "Correct / not misleading",
                "Consistent with available evidence or does not contain an evident misleading claim.",
            ],
            [
                "Stance",
                "Supportive",
                "Expresses trust, acceptance, recommendation, or encouragement toward vaccination.",
            ],
            [
                "Stance",
                "Neutral",
                "Provides information, asks a question, or reports an event without a clear orientation.",
            ],
            [
                "Stance",
                "Opposing / hesitant",
                "Expresses refusal, distrust, fear, or advice against vaccination.",
            ],
            ["Sentiment", "Positive", "Shows confidence, relief, hope, or supportive emotion."],
            ["Sentiment", "Neutral", "Does not express a clear affective polarity."],
            ["Sentiment", "Negative", "Shows fear, anger, anxiety, distrust, or blame."],
        ],
        "align": ["left", "left", "left"],
        "latex_align": "lll",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.17\textwidth}"
            r">{\raggedright\arraybackslash}p{0.22\textwidth}"
            r">{\raggedright\arraybackslash}p{0.53\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.5pt",
        "arraystretch": "1.08",
        "widths": [1450, 1900, 3250],
        "font_size": 7.4,
        "note": "Personal post-vaccination experience was not automatically labeled as misinformation unless it made an unsupported causal claim or harmful recommendation.",
    },
    "splits": {
        "caption": "Dataset split and label counts used for training and evaluation.",
        "label": "tab:splits",
        "headers": [
            "Split",
            "N",
            "Mis.",
            "Corr.",
            "Sup.",
            "Neu-S",
            "Opp.",
            "Pos.",
            "Neu-E",
            "Neg.",
        ],
        "rows": [
            ["Train", "1496", "273", "1223", "427", "622", "447", "303", "571", "622"],
            ["Validation", "167", "31", "136", "48", "66", "53", "27", "65", "75"],
            ["Gold Test", "186", "28", "158", "54", "84", "48", "40", "75", "71"],
            ["Total", "1849", "332", "1517", "529", "772", "548", "370", "711", "768"],
        ],
        "align": [
            "left",
            "center",
            "center",
            "center",
            "center",
            "center",
            "center",
            "center",
            "center",
            "center",
        ],
        "latex_align": "lrrrrrrrrr",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.14\textwidth}"
            r">{\centering\arraybackslash}p{0.075\textwidth}"
            r">{\centering\arraybackslash}p{0.065\textwidth}"
            r">{\centering\arraybackslash}p{0.085\textwidth}"
            r">{\centering\arraybackslash}p{0.065\textwidth}"
            r">{\centering\arraybackslash}p{0.105\textwidth}"
            r">{\centering\arraybackslash}p{0.065\textwidth}"
            r">{\centering\arraybackslash}p{0.065\textwidth}"
            r">{\centering\arraybackslash}p{0.105\textwidth}"
            r">{\centering\arraybackslash}p{0.065\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "1.8pt",
        "arraystretch": "1.06",
        "widths": [1000, 690, 620, 700, 610, 640, 610, 610, 640, 610],
        "font_size": 6.1,
        "cell_margins": {"top": 55, "start": 30, "bottom": 55, "end": 30},
        "nowrap_cols": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
        "note": "Mis. = misinformation; Corr. = correct/not misleading; Sup. = supportive stance; Neu-S = neutral stance; Opp. = opposing or hesitant stance; Neu-E = neutral sentiment; Pos./Neg. = positive/negative sentiment. Seven samples were dropped before modeling because word segmentation removed malformed or empty rows.",
    },
    "macro": {
        "caption": "Primary experimental results: Macro F1 on the expert-validated Gold Test Set.",
        "label": "tab:macro",
        "headers": ["Model", "Role", "Misinfo", "Stance", "Sentiment", "Mean"],
        "rows": [
            ["XLM-R-v1", "Baseline encoder", "0.7038", "0.6224", "0.6866", "0.6709"],
            ["PhoBERT-v2", "Classification engine", "0.6996", "0.6640", "0.7266", "0.6967"],
            ["Gemma-4 4B", "XAI reasoning engine", "0.6377", "0.6264", "0.7700", "0.6780"],
        ],
        "align": ["left", "left", "center", "center", "center", "center"],
        "latex_align": "llrrrr",
        "widths": [1250, 2050, 850, 850, 950, 650],
    },
    "perclass": {
        "caption": "Detailed experimental results: per-class F1 by task on the Gold Test Set.",
        "label": "tab:perclass",
        "headers": ["Task", "Class", "XLM-R-v1", "PhoBERT-v2", "Gemma-4 4B", "Support"],
        "rows": [
            ["Misinfo", "Misinformation", "0.5079", "0.5075", "0.4444", "28"],
            ["Misinfo", "Correct", "0.8997", "0.8918", "0.8309", "158"],
            ["Stance", "Supportive", "0.5495", "0.5934", "0.4528", "54"],
            ["Stance", "Opposing", "0.6387", "0.6612", "0.6905", "48"],
            ["Stance", "Neutral", "0.6790", "0.7375", "0.7360", "84"],
            ["Sentiment", "Negative", "0.7682", "0.8000", "0.8039", "71"],
            ["Sentiment", "Neutral", "0.7162", "0.7917", "0.8034", "75"],
            ["Sentiment", "Positive", "0.5753", "0.5882", "0.7027", "40"],
        ],
        "align": ["left", "left", "center", "center", "center", "center"],
        "latex_align": "llrrrr",
        "widths": [1050, 1500, 1000, 1100, 1150, 800],
        "note": "F1 is reported for each class before macro-averaging within a task.",
    },
    "calibration": {
        "caption": "Confidence calibration of PhoBERT-v2 after Temperature Scaling.",
        "label": "tab:calibration",
        "headers": [
            "Axis",
            "T",
            "ECE pre",
            "ECE post",
            "Raw conf.",
            "Cal. conf.",
            "Acc.",
            "ECE delta",
        ],
        "rows": [
            ["Misinformation", "1.82", "0.123", "0.054", "94.5%", "87.4%", "82.3%", "-56%"],
            ["Stance", "1.67", "0.198", "0.093", "86.7%", "76.3%", "67.7%", "-53%"],
            ["Sentiment", "1.35", "0.144", "0.081", "88.9%", "83.0%", "75.8%", "-44%"],
        ],
        "align": ["left", "center", "center", "center", "center", "center", "center", "center"],
        "latex_align": "lrrrrrrr",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.205\textwidth}"
            r">{\centering\arraybackslash}p{0.065\textwidth}"
            r">{\centering\arraybackslash}p{0.085\textwidth}"
            r">{\centering\arraybackslash}p{0.085\textwidth}"
            r">{\centering\arraybackslash}p{0.09\textwidth}"
            r">{\centering\arraybackslash}p{0.09\textwidth}"
            r">{\centering\arraybackslash}p{0.07\textwidth}"
            r">{\centering\arraybackslash}p{0.10\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.0pt",
        "arraystretch": "1.06",
        "widths": [1700, 620, 780, 780, 840, 840, 650, 850],
        "font_size": 6.9,
        "nowrap_cols": [0, 1, 2, 3, 4, 5, 6, 7],
        "note": "ECE = Expected Calibration Error; T = learned temperature.",
    },
    "llm": {
        "caption": "Reliability of LLM-assisted annotation against expert review.",
        "label": "tab:llm",
        "headers": ["Axis / metric", "Result", "Interpretation"],
        "rows": [
            ["Misinformation kappa", "0.0525", "Slight agreement"],
            ["Stance kappa", "0.2895", "Fair agreement"],
            ["Sentiment kappa", "0.4712", "Moderate agreement"],
            ["Mean kappa", "0.2711", "Low average agreement across axes"],
            ["Invalid structured outputs", "52/186 (28.0%)", "Requires parser and human review"],
        ],
        "align": ["left", "center", "left"],
        "latex_align": "lll",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.34\textwidth}"
            r">{\centering\arraybackslash}p{0.19\textwidth}"
            r">{\raggedright\arraybackslash}p{0.38\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.5pt",
        "arraystretch": "1.06",
        "widths": [2300, 1500, 2800],
    },
    "sentiment_stance": {
        "caption": "Cross-tabulation of sentiment and stance in the Gold Test Set.",
        "label": "tab:sentimentstance",
        "headers": ["Sentiment", "Supportive", "Opposing", "Neutral", "Total"],
        "rows": [
            ["Negative", "9", "45", "17", "71"],
            ["Neutral", "7", "3", "65", "75"],
            ["Positive", "38", "0", "2", "40"],
            ["Total", "54", "48", "84", "186"],
        ],
        "align": ["left", "center", "center", "center", "center"],
        "latex_align": "lrrrr",
        "widths": [1700, 1150, 1150, 1150, 900],
        "note": "The association was significant by chi-square test, chi-square = 189.48, df = 4, p = 6.85e-40.",
    },
    "platform_misinfo": {
        "caption": "Cross-tabulation of platform and misinformation in the Gold Test Set.",
        "label": "tab:platformmisinfo",
        "headers": ["Platform", "Misinfo", "Correct", "Total", "Misinfo rate"],
        "rows": [
            ["Facebook", "19", "58", "77", "24.7%"],
            ["YouTube", "9", "65", "74", "12.2%"],
            ["Forums / other social media", "0", "16", "16", "0.0%"],
            ["Institutional news", "0", "10", "10", "0.0%"],
            ["Academic / VFND", "0", "9", "9", "0.0%"],
            ["Total", "28", "158", "186", "15.1%"],
        ],
        "align": ["left", "center", "center", "center", "center"],
        "latex_align": "lrrrr",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.40\textwidth}"
            r">{\centering\arraybackslash}p{0.11\textwidth}"
            r">{\centering\arraybackslash}p{0.11\textwidth}"
            r">{\centering\arraybackslash}p{0.11\textwidth}"
            r">{\centering\arraybackslash}p{0.16\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.2pt",
        "arraystretch": "1.06",
        "widths": [2700, 900, 900, 850, 1250],
        "font_size": 7.2,
        "note": "The G-test was used as the main test because several expected counts were below five; G = 16.77, df = 4, p = 2.14e-3.",
    },
    "stance_misinfo": {
        "caption": "Cross-tabulation of stance and misinformation in the Gold Test Set.",
        "label": "tab:stancemisinfo",
        "headers": ["Stance", "Misinfo", "Correct", "Total", "Misinfo rate"],
        "rows": [
            ["Supportive", "1", "53", "54", "1.9%"],
            ["Opposing / hesitant", "24", "24", "48", "50.0%"],
            ["Neutral", "3", "81", "84", "3.6%"],
            ["Total", "28", "158", "186", "15.1%"],
        ],
        "align": ["left", "center", "center", "center", "center"],
        "latex_align": "lrrrr",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.36\textwidth}"
            r">{\centering\arraybackslash}p{0.12\textwidth}"
            r">{\centering\arraybackslash}p{0.12\textwidth}"
            r">{\centering\arraybackslash}p{0.12\textwidth}"
            r">{\centering\arraybackslash}p{0.17\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.2pt",
        "arraystretch": "1.06",
        "widths": [2400, 950, 950, 900, 1400],
        "font_size": 7.3,
        "note": "The association was significant by chi-square test, chi-square = 61.86, df = 2, p = 3.69e-14; the opposing-versus-supportive odds ratio was approximately 53.0.",
    },
    "hypotheses": {
        "caption": "Statistical tests for public-health risk associations.",
        "label": "tab:hypotheses",
        "headers": ["Hyp.", "Relationship", "Test", "Statistic", "p-value", "Finding"],
        "rows": [
            ["H1", "Sentiment x stance", "Chi-square, df = 4", "189.48", "6.85e-40", "Reject H0"],
            ["H2", "Platform x misinformation", "G-test, df = 4", "16.77", "2.14e-3", "Reject H0"],
            ["H3", "Stance x misinformation", "Chi-square, df = 2", "61.86", "3.69e-14", "Reject H0"],
        ],
        "align": ["center", "left", "left", "center", "center", "center"],
        "latex_align": "llllll",
        "latex_spec": (
            r"@{}>{\centering\arraybackslash}p{0.06\textwidth}"
            r">{\raggedright\arraybackslash}p{0.25\textwidth}"
            r">{\raggedright\arraybackslash}p{0.22\textwidth}"
            r">{\centering\arraybackslash}p{0.12\textwidth}"
            r">{\centering\arraybackslash}p{0.13\textwidth}"
            r">{\centering\arraybackslash}p{0.13\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.0pt",
        "arraystretch": "1.06",
        "widths": [500, 1800, 1600, 850, 900, 950],
        "font_size": 7.2,
        "note": "H0 denotes the null hypothesis of no association.",
    },
    "deployment_profile": {
        "caption": "Deployment-oriented technical profile of VaccineNLP.",
        "label": "tab:deploymentprofile",
        "headers": ["Component", "Evidence / configuration", "Operational implication"],
        "rows": [
            [
                "PhoBERT-v2 classifier",
                "~540 MB engine; 135M-parameter Vietnamese encoder; 256-token input; three task-specific heads",
                "Efficient first-tier screening for all incoming vaccine-related texts",
            ],
            [
                "Gemma-4 4B QLoRA",
                "~2.5 GB explanation engine; 4-bit NF4 QLoRA; 36.7M trainable parameters (0.46% of the base model)",
                "Selective rationale generation for low-confidence, sarcastic, implicit, or high-risk cases",
            ],
            [
                "Selective escalation",
                "Prototype design escalates only difficult cases, estimated at about 20% of inputs",
                "Reduces computation while preserving expert oversight for ambiguous content",
            ],
            [
                "Reliability controls",
                "Temperature Scaling reduced ECE by 44--56%; parser v3 reached 72.0% success (134/186)",
                "Supports confidence-aware triage and prevents free-form LLM output from acting as ground truth",
            ],
            [
                "XAI and evidence layer",
                "Integrated Gradients token attribution, calibrated confidence, consistency flags, and curated official-source anchors",
                "Provides an audit trail for public-health analysts and model reviewers",
            ],
        ],
        "align": ["left", "left", "left"],
        "latex_align": "lll",
        "latex_spec": (
            r"@{}>{\raggedright\arraybackslash}p{0.22\textwidth}"
            r">{\raggedright\arraybackslash}p{0.37\textwidth}"
            r">{\raggedright\arraybackslash}p{0.32\textwidth}@{}"
        ),
        "latex_font": "scriptsize",
        "tabcolsep": "2.0pt",
        "arraystretch": "1.08",
        "widths": [1500, 2650, 2450],
        "font_size": 7.0,
    },
}

SECTIONS = [
    {
        "title": "Introduction",
        "blocks": [
            (
                "Vaccination is one of the most effective public-health interventions, "
                "yet its population-level impact depends on public trust. During and "
                "after the COVID-19 vaccine rollout, Vietnamese digital platforms "
                "became an important arena for health communication, community debate, "
                "and misinformation [2,3,4]. The World Health Organization describes "
                "the overabundance of accurate and inaccurate information during health "
                "crises as an infodemic, a phenomenon that requires monitoring, "
                "analysis, and timely response [1,7]."
            ),
            (
                "Vaccine misinformation differs from generic fake-news detection in "
                "three important ways. First, the correctness of a claim often depends "
                "on medical evidence and public-health context rather than surface "
                "linguistic cues alone [19]. Second, vaccine debates combine factual "
                "claims, stance toward vaccination, and affective framing. Third, "
                "Vietnamese online discourse contains code-switching, teen-code "
                "abbreviations, sarcasm, and evasive expressions around vaccine-related "
                "terms. These characteristics make keyword filtering insufficient for "
                "public-health surveillance."
            ),
            (
                "The study is organized around three research questions that are "
                "retained in this proceedings paper. RQ1 asks how a "
                "Vietnamese vaccine-domain corpus can be collected, cleaned, and "
                "validated with a multi-axis annotation schema. RQ2 asks how well a "
                "Dual-Student Hybrid architecture detects misinformation while also "
                "producing interpretable explanations. RQ3 asks whether sentiment, "
                "stance, and source platform are statistically associated with "
                "misinformation risk. These questions are deliberately operational: "
                "they connect model accuracy with the needs of public-health analysts "
                "who must prioritize limited review capacity during an infodemic."
            ),
            (
                "The Vietnamese setting makes this problem especially relevant for "
                "FISAT themes in artificial intelligence, digital transformation, and "
                "applied health informatics. Vaccine discussions often appear as "
                "short comments, conversational fragments, or screenshots copied "
                "across platforms. A single comment may include a personal adverse "
                "event narrative, a rhetorical question, and an implicit recommendation "
                "against vaccination. For this reason, the system is designed not as a "
                "binary censorship tool, but as an intelligent surveillance and triage "
                "layer that surfaces uncertain or high-risk texts for expert review."
            ),
            {"subsection": "Research Contributions"},
            (
                "This paper contributes an applied intelligent-system study for the "
                "FISAT/Springer proceedings scope, at the intersection of natural "
                "language processing, trustworthy AI, health informatics, and social "
                "media analytics. Consistent with the original study protocol, the "
                "paper addresses two aims: (i) to build and evaluate Vietnamese NLP "
                "models for vaccine-related misinformation, stance, and sentiment "
                "classification; and (ii) to assess the usefulness and limitations of "
                "LLM support for medical text annotation and explanation. To make the "
                "research contribution explicit, the paper reports four concrete "
                "outputs: C1, an expert-validated Vietnamese Gold Test Set for "
                "vaccine-related misinformation, stance, and sentiment; C2, a "
                "controlled benchmark of XLM-RoBERTa, PhoBERT-v2, and Gemma-4 4B on "
                "the same evaluation split; C3, a confidence-calibrated explainable "
                "HITL workflow for trustworthy public-health deployment; and C4, "
                "empirical evidence that LLM-based annotation remains insufficient as "
                "a replacement for expert medical fact-checking."
            ),
        ],
    },
    {
        "title": "Related Work and Research Gap",
        "blocks": [
            (
                "Research on vaccine misinformation detection has progressed rapidly "
                "in English, including specialized datasets such as ANTi-Vax, where "
                "fine-tuned transformer models can reach high scores under in-domain "
                "conditions [5]. Susceptibility to COVID-19 misinformation is also "
                "shaped by social-media exposure and analytic reasoning [6]. However, "
                "direct transfer to Vietnamese vaccine discourse is limited by language "
                "shift, platform shift, and annotation-policy differences. Existing "
                "Vietnamese fake-news resources are useful but are generally small or "
                "not specific to vaccine and public-health communication."
            ),
            (
                "The full project report identifies three concrete gaps in the "
                "Vietnamese literature. First, available Vietnamese fake-news datasets "
                "are not designed around vaccine claims, so they do not capture the "
                "difference between general political misinformation and medical "
                "misinformation. Second, prior Vietnamese NLP benchmarks rarely "
                "evaluate misinformation, stance, and sentiment under the same "
                "protocol, although these axes interact in real vaccine debates. Third, "
                "there is still limited work on Vietnamese explainable AI for health "
                "communication, where practitioners need to inspect why an item was "
                "flagged before taking action."
            ),
            (
                "Transformer encoders such as BERT [10], XLM-RoBERTa [9], and "
                "PhoBERT [8], all built on the Transformer architecture [11], are well "
                "suited to text classification because they learn contextual sentence "
                "representations. PhoBERT is particularly relevant because it is "
                "pre-trained for Vietnamese and expects word-segmented input. In "
                "contrast, decoder-only LLMs are attractive for weak annotation and "
                "natural-language explanation, but their output format can be unstable "
                "and their factual judgments in medical domains require expert "
                "oversight [13,14]."
            ),
            (
                "Table 1 positions the proposed study against related work. The table "
                "is intentionally presented as a research-positioning comparison, not "
                "as a direct leaderboard, because reported scores in the literature "
                "come from different languages, platforms, annotation policies, and "
                "class distributions. Instead, it clarifies the research gap addressed "
                "by VaccineNLP: Vietnamese vaccine misinformation has not been "
                "evaluated as a combined misinformation-stance-sentiment problem with "
                "expert validation, calibration, and operational XAI in one pipeline."
            ),
            {"table": "literature_comparison"},
            (
                "Accordingly, Table 1 should be read as a research-positioning "
                "comparison rather than as a direct benchmark. The controlled "
                "experimental comparison is reported later in Tables 6 and 7, where "
                "XLM-RoBERTa, PhoBERT-v2, and Gemma-4 4B are evaluated on the same "
                "Vietnamese Gold Test Set under the same label schema."
            ),
            (
                "The distinction between encoder and decoder models is important for "
                "this study. Encoder-only models are optimized for stable "
                "representations and classification, while decoder-only LLMs produce "
                "free-form explanations that can be useful for analysts but difficult "
                "to parse into controlled labels. This motivates the Dual-Student "
                "Hybrid architecture: PhoBERT-v2 performs routine multi-task "
                "classification, and Gemma-4 4B supports explanation for difficult "
                "cases. The design aligns with a practical division of labor instead "
                "of forcing a single model to be the classifier, annotator, "
                "fact-checker, and explainer at once."
            ),
            (
                "A second gap concerns trust. Public-health AI systems must expose "
                "uncertainty and reasoning traces, not merely produce labels [22]. "
                "Calibration methods such as Temperature Scaling address the mismatch "
                "between model confidence and empirical correctness [16,17]. "
                "Explainability methods such as Integrated Gradients and natural "
                "language rationales can make model outputs more inspectable [21], "
                "but explanation does not guarantee factual validity. VaccineNLP "
                "therefore combines classification, calibration, and human-in-the-loop "
                "review rather than treating LLM output as ground truth."
            ),
        ],
    },
    {
        "title": "Materials and Methods",
        "blocks": [
            {"subsection": "Data Collection and Processing"},
            (
                "The study uses secondary public text data and a supervised "
                "machine-learning design. The unit of analysis is a Vietnamese "
                "vaccine-related post, article, video description, or comment. Data "
                "were collected from January 2020 to March 2026 from online news, "
                "Facebook, YouTube, Reddit, and public forums. Private messages were "
                "excluded. Table 2 summarizes the source-level collection and cleaning "
                "pipeline."
            ),
            (
                "Eligible texts were public Vietnamese-language items discussing "
                "human vaccination, vaccine safety, vaccine policy, adverse events, "
                "or vaccine-related trust and refusal. Exclusion criteria followed "
                "the study protocol: exact or near duplicates, spam and advertising, "
                "empty or extremely short texts, non-Vietnamese content, veterinary "
                "vaccine content, and private-channel data from platforms such as "
                "Zalo, Messenger, or Telegram. This kept the analysis focused on "
                "public discourse while reducing privacy risk."
            ),
            {"table": "sources"},
            (
                "The collection strategy was tier-based. Institutional sources such "
                "as online newspapers provided edited public-health communication; "
                "non-institutional sources such as Facebook, YouTube, Reddit, and "
                "forums provided community discussion; and academic or benchmark "
                "sources were used only as auxiliary references. The tier metadata "
                "allowed the analysis to compare platform risk while avoiding a "
                "false equivalence between curated news articles and informal social "
                "media comments."
            ),
            (
                "The corpus followed a Bronze-Silver-Gold data architecture. The "
                "Bronze layer contained raw scraped or API-derived items. The Silver "
                "layer contained cleaned text and weak labels generated by an LLM "
                "oracle. The Gold layer contained expert-reviewed labels used for "
                "independent evaluation. After deduplication and filtering, 1,856 "
                "samples were passed to annotation; 1,496 samples were used for "
                "training, 167 for validation, and 186 for the Gold Test Set."
            ),
            (
                "Vietnamese text was normalized through eight steps: HTML entity "
                "decoding, Unicode NFC normalization, URL and HTML-tag removal, "
                "hashtag splitting, emoji-to-text conversion, teen-code normalization, "
                "lowercasing, and removal of irrelevant special characters while "
                "preserving Vietnamese diacritics. Domain-specific abbreviations were "
                "expanded, for example vx to vaccine, tc to vaccination, pup to "
                "adverse reaction, and tdp to adverse effect. A relevance filter "
                "removed veterinary-vaccine and off-topic content."
            ),
            (
                "The preprocessing pipeline was designed for Vietnamese social-media "
                "noise rather than formal newspaper text alone. It retained diacritics "
                "because accent removal can collapse medically meaningful tokens, but "
                "it normalized Unicode variants and conversational abbreviations. "
                "Hashtags were split into word-like units, emojis were converted into "
                "textual sentiment cues, and URLs were removed while preserving the "
                "surrounding claim. For PhoBERT-v2, the cleaned text was subsequently "
                "word-segmented so that the model received the tokenization format "
                "expected by Vietnamese pre-training."
            ),
            {"subsection": "Annotation Schema and HITL Review"},
            (
                "Each sample was labeled on three axes. The misinformation axis is "
                "binary: misinformation versus correct or not evidently misleading "
                "information. The stance axis has three classes: supportive, neutral, "
                "and opposing or hesitant. The sentiment axis has three classes: "
                "positive, neutral, and negative. This separation between stance and "
                "sentiment follows the broader distinction between opinion orientation "
                "and affective polarity in sentiment analysis [12]. Borderline cases "
                "were adjudicated by expert review, and the expert label was treated "
                "as the reference label in the Gold Test Set."
            ),
            (
                "The annotation workflow combined LLM-assisted labeling with "
                "human-in-the-loop validation. A large LLM proposed initial labels "
                "and short rationales according to the codebook. Two reviewers with "
                "data-science expertise checked the structured output, and cases "
                "involving medical truth claims were reviewed under supervision from "
                "a public-health specialist. Disagreements and ambiguous texts were "
                "resolved through discussion, with the final consensus label used as "
                "the Gold Test reference. Personal post-vaccination experiences were "
                "not treated as misinformation unless they made unsupported causal "
                "claims or recommended harmful action."
            ),
            {"table": "schema"},
            (
                "After annotation and safety filtering, the modeling dataset contained "
                "1,849 usable samples. This is seven fewer than the 1,856 texts passed "
                "to annotation because word segmentation removed malformed or empty "
                "rows. The final split preserved the study design: 1,496 training "
                "samples, 167 validation samples, and 186 expert-validated Gold Test "
                "samples (Table 4)."
            ),
            {"table": "splits"},
            {"subsection": "Implementation Workflow"},
            (
                "The implementation can be summarized as a six-stage algorithmic "
                "workflow. Step 1 collects vaccine-related Vietnamese texts from "
                "news, social media, video comments, and public forums. Step 2 "
                "applies the Vietnamese preprocessing pipeline, including Unicode "
                "normalization, URL and HTML removal, hashtag segmentation, slang "
                "normalization, and word segmentation. Step 3 assigns three-axis "
                "labels for misinformation, stance, and sentiment through "
                "LLM-assisted annotation followed by HITL expert validation. Step 4 "
                "trains XLM-RoBERTa and PhoBERT-v2 classifiers and fine-tunes "
                "Gemma-4 4B with QLoRA for explanation-oriented reasoning. Step 5 "
                "evaluates all models on the expert-validated Gold Test Set using "
                "Macro F1, per-class F1, calibration error, parsing reliability, and "
                "agreement with expert labels. Step 6 deploys the calibrated "
                "PhoBERT-v2 output as the first-tier screening signal and escalates "
                "low-confidence or high-risk cases to the explanation and "
                "expert-review layer. This workflow is the operational algorithm used "
                "throughout the experiments, linking data preparation, modeling, "
                "evaluation, and human review in a single reproducible pipeline."
            ),
            {"subsection": "Dual-Student Hybrid Architecture"},
            (
                "VaccineNLP adopts a Dual-Student Hybrid design. The classification "
                "engine is a PhoBERT-base-v2 shared encoder with three task-specific "
                "heads for misinformation, stance, and sentiment. The baseline is "
                "XLM-RoBERTa-base trained under comparable settings. The explanation "
                "engine is a Gemma-4 4B model fine-tuned with QLoRA to generate "
                "label-aware Vietnamese rationales [15]. Figure 1 summarizes the "
                "two-tier architecture and the human-in-the-loop safety logic. The "
                "classification loss is a weighted multi-task objective:"
            ),
            {"figure": "architecture"},
            {"equation": r"L = 1.2L_{\mathrm{misinfo}} + L_{\mathrm{stance}} + L_{\mathrm{sentiment}}."},
            (
                "The misinformation loss was upweighted because misinformation is the "
                "minority but highest-risk class. In operational use, the classifier "
                "screens all incoming items and the reasoning engine supports "
                "explanation, triage, and expert review for high-risk or low-confidence "
                "cases."
            ),
            (
                "PhoBERT-v2 used a shared Vietnamese encoder with three task-specific "
                "classification heads. The training setup followed the full report: "
                "input length was capped at 256 tokens, training batches used 16 "
                "examples, validation and test batches used 32 examples, and early "
                "stopping monitored the average Macro F1 across the three axes. "
                "XLM-RoBERTa-base was trained under a comparable split and evaluation "
                "protocol as a multilingual baseline. Weighted loss was used because "
                "the misinformation label was the minority but carried the highest "
                "public-health cost."
            ),
            (
                "The encoder implementation used PhoBERT-base-v2 with a 768-dimensional "
                "pooled representation and parallel classification heads for the three "
                "annotation axes. AdamW optimization used a learning rate of 2e-5, "
                "weight decay of 0.01, three to five epochs, and a warm-up schedule of "
                "approximately 10% of training steps. Checkpoints were selected by "
                "validation loss and monitored with Macro F1 to prevent the majority "
                "classes from dominating model selection."
            ),
            (
                "Gemma-4 4B was fine-tuned with QLoRA so that a compact decoder model "
                "could generate label-aware rationales without full-parameter "
                "training. The explanation engine used longer contexts than the "
                "encoder classifier, with a maximum sequence length of 1024 tokens and "
                "an effective batch size of 16 through gradient accumulation. In the "
                "prototype, Gemma is not called for every input; it is reserved for "
                "difficult cases such as sarcasm, long conversational chains, implicit "
                "hesitation, or low calibrated confidence."
            ),
            (
                "The QLoRA configuration followed the modeling experiment: 4-bit NF4 "
                "quantization, LoRA rank r = 16, LoRA alpha = 16, and text-only target "
                "modules in the attention and MLP blocks. Only about 36.7 million "
                "parameters were trainable, approximately 0.46% of the full model. "
                "Training examples used a chat-style format in which the user message "
                "asked for vaccine misinformation analysis and the assistant response "
                "contained reasoning followed by structured labels. Loss was applied "
                "only to the response portion so the model learned to produce "
                "task-specific rationales rather than merely copy instructions."
            ),
            {"subsection": "Evaluation Protocol and Ethics"},
            (
                "Classification performance was measured by Macro F1, which is "
                "appropriate for imbalanced labels [24]. Confidence calibration was "
                "evaluated with Expected Calibration Error before and after Temperature "
                "Scaling [16]. Agreement between the LLM annotator and expert labels "
                "was measured by Cohen's kappa [18]. Associations among sentiment, "
                "stance, misinformation, and source platform were tested using "
                "chi-square or G-tests with alpha = 0.05."
            ),
            (
                "Temperature Scaling was fitted on the validation set and evaluated "
                "on the Gold Test Set. This method learns a single positive "
                "temperature that rescales logits before softmax, reducing "
                "overconfidence without changing the predicted class. For a "
                "public-health dashboard, this distinction matters: a calibrated "
                "confidence score supports review prioritization, whereas an "
                "overconfident but wrong prediction can mislead analysts."
            ),
            (
                "The statistical tests were aligned with the pre-specified hypotheses. "
                "H1 examined the association between sentiment and stance, H2 examined "
                "the association between platform and misinformation, and H3 examined "
                "whether opposing or hesitant stance concentrated misinformation more "
                "than supportive or neutral stance. The G-test was used for platform "
                "analysis because source groups were imbalanced, while chi-square "
                "tests were used for the categorical association tables."
            ),
            (
                "Error analysis was defined on randomly sampled false-positive and "
                "false-negative cases from the best-performing classifier. The "
                "error-analysis protocol grouped errors into four categories: sarcasm or irony, "
                "borderline medical content, out-of-vocabulary or emerging slang, and "
                "very short texts with insufficient context. These categories informed "
                "the future-work recommendations on sarcasm handling, active learning, "
                "and evidence-aware explanation."
            ),
            (
                "The study used only publicly available text, did not collect private "
                "messages, and did not intervene on human participants. Usernames, "
                "profile URLs, phone numbers, and other potentially identifying "
                "details were removed or anonymized before analysis."
            ),
        ],
    },
    {
        "title": "Results",
        "blocks": [
            (
                "The experimental results are organized so that each table answers a "
                "specific evaluation question. Table 5 describes the Gold Test Set "
                "used for evaluation. Table 6 is the primary model-comparison table: "
                "it reports task-level Macro F1 for XLM-RoBERTa, PhoBERT-v2, and "
                "Gemma-4 4B on the same expert-validated split. Table 7 provides the "
                "per-class F1 details behind those macro scores. Table 8 reports "
                "confidence calibration after Temperature Scaling, while Table 9 "
                "reports the reliability of LLM-assisted annotation against expert "
                "review. Tables 10--13 provide statistical evidence for public-health "
                "risk associations among sentiment, stance, platform, and "
                "misinformation."
            ),
            (
                "The Gold Test Set contains 186 expert-reviewed Vietnamese samples. "
                "Misinformation is the minority class, with 28 samples (15.1%), while "
                "158 samples were labeled as correct or not clearly misleading. For "
                "stance, 54 samples were supportive, 48 opposing or hesitant, and 84 "
                "neutral. For sentiment, 71 samples were negative, 75 neutral, and 40 "
                "positive (Table 5)."
            ),
            {"table": "labels"},
            (
                "The label distribution confirms why Macro F1 was selected as the "
                "primary metric. Accuracy alone would overstate system performance "
                "because the correct/not-misleading class dominates the "
                "misinformation axis. A classifier that misses most harmful "
                "misinformation could still appear acceptable under accuracy, but "
                "would be unsafe for infodemic surveillance. Macro F1 gives each class "
                "equal weight and therefore makes minority-class weakness visible."
            ),
            (
                "Table 6 is therefore the main experimental-results table for model "
                "performance. It summarizes Macro F1 across the three tasks. "
                "PhoBERT-v2 "
                "achieved the highest average Macro F1 (0.6967), suggesting that a "
                "Vietnamese-specific encoder is effective for multi-task health-text "
                "classification. XLM-RoBERTa slightly outperformed PhoBERT-v2 on the "
                "misinformation axis (0.7038 versus 0.6996), while Gemma-4 4B "
                "performed best on sentiment (0.7700), consistent with the strength "
                "of generative models in affective and contextual interpretation."
            ),
            {"table": "macro"},
            (
                "The three models therefore play complementary roles. PhoBERT-v2 is "
                "the most balanced model across tasks, with a smaller performance "
                "range between misinformation, stance, and sentiment. XLM-RoBERTa "
                "remains competitive for misinformation, possibly because its "
                "multilingual pre-training captures broader vaccine discourse. "
                "Gemma-4 4B is strongest for sentiment, but its weaker structured "
                "classification performance and parsing instability make it better "
                "suited as an explanation engine than as the sole decision engine."
            ),
            (
                "The specialization pattern is also visible in the between-task "
                "Macro F1 range. The spread between the best and worst task was "
                "0.0626 for PhoBERT-v2, 0.0814 for XLM-RoBERTa, and 0.1436 for "
                "Gemma-4 4B. This narrower range supports PhoBERT-v2 as the stable "
                "screening engine, while the wider Gemma-4 4B range supports its "
                "selective use for explanation and escalation rather than universal "
                "classification."
            ),
            (
                "Per-class results show that misinformation detection remains the "
                "hardest operational task (Table 7). For the minority misinformation "
                "class, F1 was 0.5079 for XLM-RoBERTa, 0.5075 for PhoBERT-v2, and "
                "0.4444 for Gemma-4 4B. In contrast, the correct-information class "
                "reached F1 scores close to 0.89 for the encoder models. This "
                "asymmetry reflects both label imbalance and the linguistic camouflage "
                "used in anti-vaccine content."
            ),
            {"table": "perclass"},
            {"figure": "per_class_f1"},
            (
                "The stance and sentiment axes show a different error profile. Neutral "
                "stance and neutral sentiment are common in news-style reporting and "
                "informational comments, while opposing stance often co-occurs with "
                "fear, distrust, or blame. Gemma-4 4B achieved strong scores on "
                "negative and neutral sentiment, which is consistent with its ability "
                "to interpret affective language, but it was less reliable on "
                "supportive stance. This pattern reinforces the value of evaluating "
                "the three axes separately rather than reporting only an aggregate "
                "score."
            ),
            (
                "Temperature Scaling substantially improved confidence reliability "
                "(Table 8). For misinformation, ECE decreased from 0.123 to 0.054; "
                "for stance, from 0.198 to 0.093; and for sentiment, from 0.144 to "
                "0.081. These reductions of 44--56% are important for deployment "
                "because public-health analysts need calibrated risk scores, not only "
                "predicted labels."
            ),
            {"table": "calibration"},
            {"figure": "calibration"},
            (
                "The calibration curves also reveal that stance is the most difficult "
                "axis to calibrate. Even after Temperature Scaling, stance ECE remains "
                "0.093, higher than misinformation and sentiment. In practice, this "
                "means stance predictions should be displayed with explicit confidence "
                "and should trigger human review when confidence is low or when the "
                "text mixes supportive and hesitant cues."
            ),
            (
                "The LLM-assisted annotation analysis exposed a major safety issue. "
                "Agreement between the LLM annotator and expert labels was very low "
                "for misinformation, fair for stance, and moderate for sentiment "
                "(Table 9). In addition, 28.0% of Gemma-4 4B outputs on the Gold Test "
                "Set could not be parsed into the required structure. The result "
                "supports a human-in-the-loop workflow in which LLMs propose labels "
                "and rationales, but expert review remains mandatory for medical truth "
                "claims."
            ),
            {"table": "llm"},
            (
                "Parser analysis from the system evaluation shows why this matters for "
                "deployment. The improved parser handled multi-block reasoning, "
                "negation, contrastive conjunctions such as 'but' and 'however', and "
                "partial outputs when at least two of three labels were recoverable. "
                "Even with these safeguards, 52 of 186 outputs were invalid. The "
                "system therefore treats LLM explanations as auxiliary evidence and "
                "does not let a free-form rationale override expert-validated labels."
            ),
            (
                "The public-health association tests were statistically significant "
                "and the cross-tabulations show where the risk concentrates. First, "
                "negative sentiment was strongly associated with anti-vaccine stance "
                "(Table 10). Among 48 opposing or hesitant samples, 45 were negative "
                "(93.8%) and none were positive. By contrast, 38 of 54 supportive "
                "samples were positive. This pattern indicates that anti-vaccine "
                "stance is often packaged through fear, anger, distrust, or blame."
            ),
            {"table": "sentiment_stance"},
            (
                "Second, platform distribution was associated with misinformation "
                "(Table 11). Facebook accounted for 19 misinformation samples among "
                "77 Facebook samples (24.7%), while YouTube accounted for 9 among 74 "
                "(12.2%). The other source groups had no misinformation samples in "
                "the Gold Test Set. This does not imply that institutional media are "
                "perfect, but it does show that editorial control and public-source "
                "credibility are meaningful features for risk triage."
            ),
            {"table": "platform_misinfo"},
            (
                "Third, opposing or hesitant stance was highly concentrated with "
                "misinformation (Table 12). Twenty-four of 48 opposing samples "
                "(50.0%) contained misinformation, compared with only 1 of 54 "
                "supportive samples (1.9%) and 3 of 84 neutral samples (3.6%). The "
                "opposing-versus-supportive odds ratio was approximately 53.0, a very "
                "large effect size for an applied public-health signal."
            ),
            {"table": "stance_misinfo"},
            (
                "Table 13 summarizes the hypothesis tests. H1 and H3 were significant "
                "at extremely small p-values, while H2 remained significant under the "
                "G-test despite the small cells in several platform groups. The "
                "results support a triage design in which negative affect, opposing "
                "stance, and high-risk platform context increase the priority of "
                "expert review."
            ),
            {"table": "hypotheses"},
            (
                "Together, the association tests support the project's "
                "misinformation-propagation triangle: negative affect, opposing or "
                "hesitant stance, and high-risk social platforms tend to reinforce one "
                "another. The finding does not imply that every negative comment is "
                "misinformation, nor that all Facebook or YouTube content is harmful. "
                "Instead, it gives analysts a statistically grounded triage signal for "
                "deciding which public posts should be reviewed first."
            ),
        ],
    },
    {
        "title": "Discussion",
        "blocks": [
            (
                "The results indicate that Vietnamese vaccine misinformation "
                "surveillance benefits from task-specialized model roles. A "
                "Vietnamese encoder is the most stable backbone for routine "
                "classification, while a generative reasoning model is useful for "
                "interpretable explanations and sentiment-rich cases. The proposed "
                "Dual-Student Hybrid design follows a frugal and trustworthy-AI "
                "principle: use the efficient classifier for all inputs and reserve "
                "the more expensive reasoning model for cases that require explanation "
                "or human review."
            ),
            (
                "This specialization is consistent with the technical nature of the "
                "tasks. Misinformation classification often requires recognizing "
                "unsupported causality, distorted evidence, or claims that are true in "
                "one context but misleading in another. Sentiment analysis, by "
                "contrast, depends more heavily on affective wording and pragmatic "
                "tone. Stance lies between these tasks because a user may express fear "
                "while still accepting vaccination, or may report an adverse event "
                "without recommending refusal. A multi-task framework is therefore "
                "more faithful to public-health communication than a single "
                "misinformation label."
            ),
            (
                "Calibration should be interpreted as part of the system design, not "
                "as a post-hoc cosmetic step. In a review dashboard, the confidence "
                "score decides whether a post is routed to routine monitoring, expert "
                "review, or evidence retrieval. If the model is overconfident on "
                "ambiguous anti-vaccine content, analysts may under-review risky "
                "items. If it is overconfident on legitimate concerns, the system may "
                "discourage open discussion and reduce trust. Temperature Scaling "
                "therefore contributes to responsible deployment by aligning model "
                "confidence with observed correctness."
            ),
            (
                "The implemented prototype also demonstrates how the model outputs can "
                "be exposed to public-health analysts as a review queue rather than as "
                "automatic medical truth judgments. Figure 4 illustrates the "
                "conversation-flow analysis screen, where comments are triaged with "
                "confidence signals and evidence-oriented explanations."
            ),
            {"figure": "prototype"},
            (
                "The prototype follows a client-server design with a web interface, an "
                "API analysis layer, model services, and persistent storage for history "
                "and review queues. The operational modules mirror the implemented system: "
                "single-text multi-task analysis, conversation-flow analysis, campaign "
                "or narrative detection, batch processing, analysis history, manual "
                "review, and system health monitoring. The classifier returns labels "
                "and softmax distributions for the three axes, while the explanation "
                "layer can add reasoning, salient text spans, and evidence status."
            ),
            (
                "From a deployment perspective, the important design decision is "
                "selective escalation. Clear low-risk cases can be handled by the "
                "PhoBERT-v2 classifier, while ambiguous, sarcastic, low-confidence, or "
                "high-risk cases are escalated to the reasoning layer and then to human "
                "review. The prototype also records a consistency flag such as "
                "plausible, unusual, high-risk, or evasion-suspected. This makes the "
                "system closer to an analyst workbench than a simple demo classifier."
            ),
            (
                "The proposed evidence layer combines retrieval-augmented generation "
                "with a curated public-health knowledge base. In the prototype, "
                "official sources such as WHO, CDC, and Ministry of Health guidance "
                "were treated as preferred anchors, while scientific and news sources "
                "supported broader context. This approach is important because an "
                "LLM-generated explanation can be fluent but still medically unsafe if "
                "it is not grounded in evidence."
            ),
            (
                "Table 14 summarizes the deployment-oriented technical profile "
                "extracted from the project implementation. The key point is that the "
                "system was not designed as a single heavy model. It combines a "
                "lightweight Vietnamese classifier, a parameter-efficient explanation "
                "engine, calibration, structured parsing, and human review into an "
                "operational workflow for public-health monitoring."
            ),
            {"table": "deployment_profile"},
            (
                "The low LLM-expert agreement on misinformation is a central finding. "
                "It cautions against using LLM-generated labels as medical ground "
                "truth without validation, especially when the task requires "
                "distinguishing personal experience, outdated information, causal "
                "misinterpretation, and unsupported medical claims. In a public-health "
                "workflow, a false positive can unfairly flag legitimate concerns, "
                "while a false negative can let harmful misinformation spread through "
                "online diffusion dynamics [20]."
            ),
            (
                "The system also implemented token-level explanation with Integrated "
                "Gradients for PhoBERT-v2. This complements Gemma-generated rationales "
                "because it exposes which tokens contributed to the classifier output. "
                "For analysts, natural-language reasoning is easier to read; for model "
                "auditors, attribution scores are more reproducible and less prone to "
                "hallucinated justification. The combined XAI strategy is therefore "
                "not merely decorative. It gives different stakeholders different "
                "forms of inspectability."
            ),
            (
                "The association results also have practical value. The combination "
                "of negative sentiment, opposing stance, and social-platform context "
                "forms a useful risk triage signal for infodemic management. Rather "
                "than treating all vaccine discussions as equal, health agencies can "
                "prioritize review queues for content where these signals converge, "
                "while still preserving expert judgment and avoiding automatic "
                "censorship-like decisions."
            ),
            (
                "For Vietnamese public-health agencies, the most realistic use case is "
                "not full automation but early warning. A local CDC or digital-health "
                "team could run the classifier over public comments, rank items by "
                "misinformation risk and calibrated uncertainty, and route selected "
                "items to expert review. The system output could then inform targeted "
                "communication, FAQ updates, or community engagement. This workflow "
                "keeps accountability with trained professionals while using NLP to "
                "reduce the manual burden of monitoring thousands of posts."
            ),
        ],
    },
    {
        "title": "Limitations and Future Work",
        "blocks": [
            (
                "The main limitation is the size and distribution of the Gold Test "
                "Set. With only 28 misinformation samples, minority-class F1 has a "
                "wide uncertainty range. The corpus is also dominated by Facebook and "
                "YouTube, so generalization to other platforms requires further "
                "validation. Engagement metrics such as shares, reactions, and view "
                "counts were not consistently retained after schema normalization, "
                "preventing a direct analysis of viral spread."
            ),
            (
                "A second limitation is that the dataset is Vietnamese-specific and "
                "reflects the language practices of selected platforms during "
                "2020--2026. The models may not generalize to future vaccine topics, "
                "new slang, private messaging ecosystems, or regional dialects without "
                "further annotation. The study also evaluates text-level "
                "classification; it does not reconstruct user networks, diffusion "
                "paths, or cross-platform reposting chains. These missing signals are "
                "important for measuring how misinformation actually spreads."
            ),
            (
                "A third limitation concerns LLM use. Gemma-4 4B is helpful for "
                "explanations, but the parse-failure rate and low kappa on medical "
                "truth claims show that free-form LLM output is not sufficiently "
                "reliable as a labeling authority. The study therefore treats LLM "
                "labels as weak supervision and auxiliary reasoning rather than as "
                "ground truth. Future versions should enforce stricter structured "
                "decoding, run adversarial prompt tests, and compare multiple LLM "
                "families under the same expert-validated protocol."
            ),
            (
                "Future work should expand expert annotation, include regional "
                "Southeast Asian multilingual data, and evaluate retrieval-augmented "
                "generation against official sources such as WHO and Vietnamese "
                "Ministry of Health guidance [23]. Further calibration methods such "
                "as Beta Calibration or isotonic regression should be compared with "
                "Temperature Scaling, particularly for the stance axis where ECE "
                "remained 0.093 after calibration [17]."
            ),
            (
                "The next technical step is an active-learning loop. The system can "
                "select uncertain, high-impact, or disagreement-heavy samples for "
                "expert review, then periodically update the classifier and "
                "calibration layer. This would make the Gold Set grow in the areas "
                "where the model is weakest rather than sampling new data uniformly. "
                "Another direction is evidence-aware generation: rationales should be "
                "grounded in retrieved documents from WHO, CDC, the Vietnamese "
                "Ministry of Health, and peer-reviewed biomedical sources, with clear "
                "separation between model reasoning and cited evidence."
            ),
        ],
    },
    {
        "title": "Conclusion",
        "blocks": [
            (
                "VaccineNLP demonstrates that multi-task Vietnamese NLP can support "
                "public-health infodemic surveillance, but only under a cautious "
                "human-in-the-loop design. PhoBERT-v2 achieved the strongest average "
                "classification performance, Temperature Scaling improved confidence "
                "reliability, and statistical tests identified meaningful links "
                "between sentiment, stance, platform, and misinformation risk. At the "
                "same time, very low LLM-expert agreement for misinformation shows "
                "that expert validation remains essential. The system is best "
                "understood as a decision-support and triage layer for public-health "
                "analysts, not an autonomous medical fact-checker."
            ),
            (
                "By integrating Vietnamese-specific preprocessing, multi-axis "
                "annotation, calibrated transformer classification, LLM-assisted "
                "explanation, and statistical public-health analysis, the study "
                "provides a reproducible foundation for future Vietnamese medical NLP "
                "systems. The central lesson is practical: trustworthy vaccine "
                "misinformation surveillance requires both intelligent automation and "
                "human accountability."
            )
        ],
    },
]

REFERENCES = [
    "World Health Organization. Managing the COVID-19 infodemic: promoting healthy behaviours and mitigating the harm from misinformation and disinformation. WHO (2020).",
    "Duong, M.C., Nguyen, H.T., Duong, M.: Evaluating COVID-19 vaccine hesitancy: a qualitative study from Vietnam. Diabetes Metab. Syndr. 16(1), 102363 (2022). https://doi.org/10.1016/j.dsx.2021.102363",
    "Ha, N., Chi, N., Van Nuil, J., Thwaites, L., Chambers, M.: Media portrayal of vaccine: a content analysis of Vietnam online news about a pentavalent vaccine in the Expanded Program of Immunization. Wellcome Open Res. 7, 271 (2022). https://doi.org/10.12688/wellcomeopenres.18457.1",
    "Bui, N., Nguyen, H., Duong, B., et al.: Lessons learned from Vietnam's first COVID-19 vaccine rollout: tackling vaccine hesitancy and misinformation for future pandemic responses. Front. Public Health 13, 1633756 (2025). https://doi.org/10.3389/fpubh.2025.1633756",
    "Hayawi, K., Shahriar, S., Serhani, M.A., Taleb, I., Mathew, S.S.: ANTi-Vax: a novel Twitter dataset for COVID-19 vaccine misinformation detection. Public Health 203, 23--30 (2022). https://doi.org/10.1016/j.puhe.2021.11.022",
    "Roozenbeek, J., Schneider, C.R., Dryhurst, S., et al.: Susceptibility to misinformation about COVID-19 around the world. R. Soc. Open Sci. 7(10), 201199 (2020). https://doi.org/10.1098/rsos.201199",
    "Eysenbach, G.: Infodemiology: the epidemiology of (mis)information. Am. J. Med. 113(9), 763--765 (2002). https://doi.org/10.1016/S0002-9343(02)01473-0",
    "Nguyen, D.Q., Nguyen, A.T.: PhoBERT: Pre-trained language models for Vietnamese. In: Findings of EMNLP, pp. 1037--1042 (2020). https://doi.org/10.18653/v1/2020.findings-emnlp.92",
    "Conneau, A., Khandelwal, K., Goyal, N., et al.: Unsupervised cross-lingual representation learning at scale. In: ACL, pp. 8440--8451 (2020). https://doi.org/10.18653/v1/2020.acl-main.747",
    "Devlin, J., Chang, M.W., Lee, K., Toutanova, K.: BERT: pre-training of deep bidirectional transformers for language understanding. In: NAACL-HLT, pp. 4171--4186 (2019). https://doi.org/10.18653/v1/N19-1423",
    "Vaswani, A., Shazeer, N., Parmar, N., et al.: Attention is all you need. In: NeurIPS, pp. 5998--6008 (2017).",
    "Pang, B., Lee, L.: Opinion mining and sentiment analysis. Found. Trends Inf. Retr. 2(1--2), 1--135 (2008). https://doi.org/10.1561/1500000011",
    "Gilardi, F., Alizadeh, M., Kubli, M.: ChatGPT outperforms crowd workers for text-annotation tasks. Proc. Natl. Acad. Sci. 120(30), e2305016120 (2023). https://doi.org/10.1073/pnas.2305016120",
    "Singhal, K., Azizi, S., Tu, T., et al.: Large language models encode clinical knowledge. Nature 620, 172--180 (2023). https://doi.org/10.1038/s41586-023-06291-2",
    "Dettmers, T., Pagnoni, A., Holtzman, A., Zettlemoyer, L.: QLoRA: efficient finetuning of quantized LLMs. In: NeurIPS (2023).",
    "Guo, C., Pleiss, G., Sun, Y., Weinberger, K.Q.: On calibration of modern neural networks. In: ICML, pp. 1321--1330 (2017).",
    "Desai, S., Durrett, G.: Calibration of pre-trained transformers. In: EMNLP, pp. 295--302 (2020). https://doi.org/10.18653/v1/2020.emnlp-main.21",
    "Landis, J.R., Koch, G.G.: The measurement of observer agreement for categorical data. Biometrics 33(1), 159--174 (1977). https://doi.org/10.2307/2529310",
    "Vraga, E.K., Bode, L.: Defining misinformation and understanding its bounded nature. Political Communication 37(1), 136--144 (2020). https://doi.org/10.1080/10584609.2020.1716500",
    "Vosoughi, S., Roy, D., Aral, S.: The spread of true and false news online. Science 359(6380), 1146--1151 (2018). https://doi.org/10.1126/science.aap9559",
    "Sundararajan, M., Taly, A., Yan, Q.: Axiomatic attribution for deep networks. In: ICML, pp. 3319--3328 (2017).",
    "Holzinger, A., Langs, G., Denk, H., Zatloukal, K., Muller, H.: Causability and explainability of artificial intelligence in medicine. WIREs Data Min. Knowl. Discov. 9(4), e1312 (2019). https://doi.org/10.1002/widm.1312",
    "Lewis, P., Perez, E., Piktus, A., et al.: Retrieval-augmented generation for knowledge-intensive NLP tasks. In: NeurIPS, pp. 9459--9474 (2020).",
    "He, H., Garcia, E.A.: Learning from imbalanced data. IEEE Trans. Knowl. Data Eng. 21(9), 1263--1284 (2009). https://doi.org/10.1109/TKDE.2008.239",
    "Mu, Y., Jin, M., Grimshaw, C., Scarton, C., Bontcheva, K., Song, X.: VaxxHesitancy: a dataset for studying hesitancy towards COVID-19 vaccination on Twitter. arXiv:2301.06660 (2023).",
    "Dadgar, S.M.H., Ghatee, M.: Checkovid: a COVID-19 misinformation detection system on Twitter using network and content mining perspectives. arXiv:2107.09768 (2021).",
    "Joshi, G., Srivastava, A., Yagnik, B., Hasan, M., Saiyed, Z., Gabralla, L.A., Abraham, A., Walambe, R., Kotecha, K.: Explainable misinformation detection across multiple social media platforms. arXiv:2203.11724 (2022).",
]

BIB_KEYS = [
    "who2020infodemic",
    "duong2022hesitancy",
    "ha2022media",
    "bui2025vaccine",
    "hayawi2022antivax",
    "roozenbeek2020susceptibility",
    "eysenbach2002infodemiology",
    "nguyen2020phobert",
    "conneau2020xlmr",
    "devlin2019bert",
    "vaswani2017attention",
    "pang2008opinion",
    "gilardi2023chatgpt",
    "singhal2023medpalm",
    "dettmers2023qlora",
    "guo2017calibration",
    "desai2020calibration",
    "landis1977kappa",
    "vraga2020misinformation",
    "vosoughi2018spread",
    "sundararajan2017integrated",
    "holzinger2019causability",
    "lewis2020rag",
    "he2009imbalanced",
    "mu2023vaxxhesitancy",
    "dadgar2021checkovid",
    "joshi2022explainable",
]

BIBTEX_ENTRIES = [
    r"""@misc{who2020infodemic,
  author       = {{World Health Organization}},
  title        = {Managing the {COVID-19} infodemic: promoting healthy behaviours and mitigating the harm from misinformation and disinformation},
  year         = {2020},
  institution  = {World Health Organization}
}""",
    r"""@article{duong2022hesitancy,
  author  = {Duong, M. C. and Nguyen, H. T. and Duong, M.},
  title   = {Evaluating {COVID-19} vaccine hesitancy: a qualitative study from {Vietnam}},
  journal = {Diabetes and Metabolic Syndrome},
  volume  = {16},
  number  = {1},
  pages   = {102363},
  year    = {2022},
  doi     = {10.1016/j.dsx.2021.102363}
}""",
    r"""@article{ha2022media,
  author  = {Ha, N. and Chi, N. and Van Nuil, J. and Thwaites, L. and Chambers, M.},
  title   = {Media portrayal of vaccine: a content analysis of {Vietnam} online news about a pentavalent vaccine in the Expanded Program of Immunization},
  journal = {Wellcome Open Research},
  volume  = {7},
  pages   = {271},
  year    = {2022},
  doi     = {10.12688/wellcomeopenres.18457.1}
}""",
    r"""@article{bui2025vaccine,
  author  = {Bui, N. and Nguyen, H. and Duong, B. and others},
  title   = {Lessons learned from {Vietnam}'s first {COVID-19} vaccine rollout: tackling vaccine hesitancy and misinformation for future pandemic responses},
  journal = {Frontiers in Public Health},
  volume  = {13},
  pages   = {1633756},
  year    = {2025},
  doi     = {10.3389/fpubh.2025.1633756}
}""",
    r"""@article{hayawi2022antivax,
  author  = {Hayawi, K. and Shahriar, S. and Serhani, M. A. and Taleb, I. and Mathew, S. S.},
  title   = {{ANTi-Vax}: a novel Twitter dataset for {COVID-19} vaccine misinformation detection},
  journal = {Public Health},
  volume  = {203},
  pages   = {23--30},
  year    = {2022},
  doi     = {10.1016/j.puhe.2021.11.022}
}""",
    r"""@article{roozenbeek2020susceptibility,
  author  = {Roozenbeek, J. and Schneider, C. R. and Dryhurst, S. and others},
  title   = {Susceptibility to misinformation about {COVID-19} around the world},
  journal = {Royal Society Open Science},
  volume  = {7},
  number  = {10},
  pages   = {201199},
  year    = {2020},
  doi     = {10.1098/rsos.201199}
}""",
    r"""@article{eysenbach2002infodemiology,
  author  = {Eysenbach, G.},
  title   = {Infodemiology: the epidemiology of (mis)information},
  journal = {American Journal of Medicine},
  volume  = {113},
  number  = {9},
  pages   = {763--765},
  year    = {2002},
  doi     = {10.1016/S0002-9343(02)01473-0}
}""",
    r"""@inproceedings{nguyen2020phobert,
  author    = {Nguyen, D. Q. and Nguyen, A. T.},
  title     = {{PhoBERT}: Pre-trained language models for {Vietnamese}},
  booktitle = {Findings of EMNLP},
  pages     = {1037--1042},
  year      = {2020},
  doi       = {10.18653/v1/2020.findings-emnlp.92}
}""",
    r"""@inproceedings{conneau2020xlmr,
  author    = {Conneau, A. and Khandelwal, K. and Goyal, N. and others},
  title     = {Unsupervised cross-lingual representation learning at scale},
  booktitle = {Proceedings of ACL},
  pages     = {8440--8451},
  year      = {2020},
  doi       = {10.18653/v1/2020.acl-main.747}
}""",
    r"""@inproceedings{devlin2019bert,
  author    = {Devlin, J. and Chang, M. W. and Lee, K. and Toutanova, K.},
  title     = {{BERT}: pre-training of deep bidirectional transformers for language understanding},
  booktitle = {Proceedings of NAACL-HLT},
  pages     = {4171--4186},
  year      = {2019},
  doi       = {10.18653/v1/N19-1423}
}""",
    r"""@inproceedings{vaswani2017attention,
  author    = {Vaswani, A. and Shazeer, N. and Parmar, N. and others},
  title     = {Attention is all you need},
  booktitle = {Advances in Neural Information Processing Systems},
  pages     = {5998--6008},
  year      = {2017}
}""",
    r"""@article{pang2008opinion,
  author  = {Pang, B. and Lee, L.},
  title   = {Opinion mining and sentiment analysis},
  journal = {Foundations and Trends in Information Retrieval},
  volume  = {2},
  number  = {1--2},
  pages   = {1--135},
  year    = {2008},
  doi     = {10.1561/1500000011}
}""",
    r"""@article{gilardi2023chatgpt,
  author  = {Gilardi, F. and Alizadeh, M. and Kubli, M.},
  title   = {{ChatGPT} outperforms crowd workers for text-annotation tasks},
  journal = {Proceedings of the National Academy of Sciences},
  volume  = {120},
  number  = {30},
  pages   = {e2305016120},
  year    = {2023},
  doi     = {10.1073/pnas.2305016120}
}""",
    r"""@article{singhal2023medpalm,
  author  = {Singhal, K. and Azizi, S. and Tu, T. and others},
  title   = {Large language models encode clinical knowledge},
  journal = {Nature},
  volume  = {620},
  pages   = {172--180},
  year    = {2023},
  doi     = {10.1038/s41586-023-06291-2}
}""",
    r"""@inproceedings{dettmers2023qlora,
  author    = {Dettmers, T. and Pagnoni, A. and Holtzman, A. and Zettlemoyer, L.},
  title     = {{QLoRA}: efficient finetuning of quantized {LLMs}},
  booktitle = {Advances in Neural Information Processing Systems},
  year      = {2023}
}""",
    r"""@inproceedings{guo2017calibration,
  author    = {Guo, C. and Pleiss, G. and Sun, Y. and Weinberger, K. Q.},
  title     = {On calibration of modern neural networks},
  booktitle = {Proceedings of ICML},
  pages     = {1321--1330},
  year      = {2017}
}""",
    r"""@inproceedings{desai2020calibration,
  author    = {Desai, S. and Durrett, G.},
  title     = {Calibration of pre-trained transformers},
  booktitle = {Proceedings of EMNLP},
  pages     = {295--302},
  year      = {2020},
  doi       = {10.18653/v1/2020.emnlp-main.21}
}""",
    r"""@article{landis1977kappa,
  author  = {Landis, J. R. and Koch, G. G.},
  title   = {The measurement of observer agreement for categorical data},
  journal = {Biometrics},
  volume  = {33},
  number  = {1},
  pages   = {159--174},
  year    = {1977},
  doi     = {10.2307/2529310}
}""",
    r"""@article{vraga2020misinformation,
  author  = {Vraga, E. K. and Bode, L.},
  title   = {Defining misinformation and understanding its bounded nature},
  journal = {Political Communication},
  volume  = {37},
  number  = {1},
  pages   = {136--144},
  year    = {2020},
  doi     = {10.1080/10584609.2020.1716500}
}""",
    r"""@article{vosoughi2018spread,
  author  = {Vosoughi, S. and Roy, D. and Aral, S.},
  title   = {The spread of true and false news online},
  journal = {Science},
  volume  = {359},
  number  = {6380},
  pages   = {1146--1151},
  year    = {2018},
  doi     = {10.1126/science.aap9559}
}""",
    r"""@inproceedings{sundararajan2017integrated,
  author    = {Sundararajan, M. and Taly, A. and Yan, Q.},
  title     = {Axiomatic attribution for deep networks},
  booktitle = {Proceedings of ICML},
  pages     = {3319--3328},
  year      = {2017}
}""",
    r"""@article{holzinger2019causability,
  author  = {Holzinger, A. and Langs, G. and Denk, H. and Zatloukal, K. and Muller, H.},
  title   = {Causability and explainability of artificial intelligence in medicine},
  journal = {WIREs Data Mining and Knowledge Discovery},
  volume  = {9},
  number  = {4},
  pages   = {e1312},
  year    = {2019},
  doi     = {10.1002/widm.1312}
}""",
    r"""@inproceedings{lewis2020rag,
  author    = {Lewis, P. and Perez, E. and Piktus, A. and others},
  title     = {Retrieval-augmented generation for knowledge-intensive {NLP} tasks},
  booktitle = {Advances in Neural Information Processing Systems},
  pages     = {9459--9474},
  year      = {2020}
}""",
    r"""@article{he2009imbalanced,
  author  = {He, H. and Garcia, E. A.},
  title   = {Learning from imbalanced data},
  journal = {IEEE Transactions on Knowledge and Data Engineering},
  volume  = {21},
  number  = {9},
  pages   = {1263--1284},
  year    = {2009},
  doi     = {10.1109/TKDE.2008.239}
}""",
    r"""@misc{mu2023vaxxhesitancy,
  author       = {Mu, Y. and Jin, M. and Grimshaw, C. and Scarton, C. and Bontcheva, K. and Song, X.},
  title        = {{VaxxHesitancy}: a dataset for studying hesitancy towards {COVID-19} vaccination on {Twitter}},
  year         = {2023},
  eprint       = {2301.06660},
  archivePrefix = {arXiv}
}""",
    r"""@misc{dadgar2021checkovid,
  author       = {Dadgar, S. M. H. and Ghatee, M.},
  title        = {{Checkovid}: a {COVID-19} misinformation detection system on {Twitter} using network and content mining perspectives},
  year         = {2021},
  eprint       = {2107.09768},
  archivePrefix = {arXiv}
}""",
    r"""@misc{joshi2022explainable,
  author       = {Joshi, G. and Srivastava, A. and Yagnik, B. and Hasan, M. and Saiyed, Z. and Gabralla, L. A. and Abraham, A. and Walambe, R. and Kotecha, K.},
  title        = {Explainable misinformation detection across multiple social media platforms},
  year         = {2022},
  eprint       = {2203.11724},
  archivePrefix = {arXiv}
}""",
]

if len(BIB_KEYS) != len(REFERENCES) or len(BIBTEX_ENTRIES) != len(REFERENCES):
    raise ValueError("Reference, BibTeX key, and BibTeX entry counts must match.")

CITE_KEYS = {str(idx): key for idx, key in enumerate(BIB_KEYS, 1)}


def latex_escape(text):
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in str(text))


def latex_text(text):
    escaped = latex_escape(text)

    def replace_citation(match):
        keys = []
        for number in match.group(1).split(","):
            number = number.strip()
            if number in CITE_KEYS:
                keys.append(CITE_KEYS[number])
        if not keys:
            return match.group(0)
        return r"\cite{" + ",".join(keys) + "}"

    return re.sub(r"\[([0-9]+(?:\s*,\s*[0-9]+)*)\]", replace_citation, escaped)


def build_bib():
    return "\n\n".join(BIBTEX_ENTRIES) + "\n"


def build_latex_table(key):
    table = TABLES[key]
    latex_spec = table.get("latex_spec", table["latex_align"])
    lines = [
        r"\begin{table}[t]",
        rf"\caption{{{latex_escape(table['caption'])}}}\label{{{table['label']}}}",
        r"\centering",
        rf"\{table.get('latex_font', 'small')}",
        rf"\setlength{{\tabcolsep}}{{{table.get('tabcolsep', '4pt')}}}",
        rf"\renewcommand{{\arraystretch}}{{{table.get('arraystretch', '1.05')}}}",
        rf"\begin{{tabular}}{{{latex_spec}}}",
        r"\hline",
        " & ".join(latex_text(h) for h in table["headers"]) + r"\\",
        r"\hline",
    ]
    for row in table["rows"]:
        lines.append(" & ".join(latex_text(v) for v in row) + r"\\")
    lines.append(r"\hline")
    lines.append(r"\end{tabular}")
    if table.get("note"):
        lines.append(r"\vspace{2pt}")
        lines.append(
            r"\parbox{\textwidth}{\footnotesize\emph{Note.} "
            + latex_text(table["note"])
            + "}"
        )
    lines.append(r"\end{table}")
    return "\n".join(lines)


def build_latex_figure(key):
    figure = FIGURES[key]
    lines = [
        r"\begin{figure}[t]",
        r"\centering",
        rf"\includegraphics[width={figure['latex_width']}]{{figures/{figure['filename']}}}",
        rf"\caption{{{latex_escape(figure['caption'])}}}\label{{{figure['label']}}}",
        r"\end{figure}",
    ]
    return "\n".join(lines)


def build_latex():
    author_tex = " \\and ".join(
        f"{latex_escape(author['name'])}\\inst{{{author['inst']}}}" for author in AUTHORS
    )
    institute_tex = " \\and ".join(latex_escape(institution) for institution in INSTITUTIONS)
    lines = [
        r"\documentclass[runningheads]{llncs}",
        r"\usepackage{graphicx}",
        r"\usepackage{array}",
        r"\usepackage{url}",
        "",
        r"\begin{document}",
        "",
        rf"\title{{{latex_escape(TITLE)}}}",
        rf"\titlerunning{{{latex_escape(TITLE_RUNNING)}}}",
        "",
        rf"\author{{{author_tex}}}",
        rf"\authorrunning{{{latex_escape(AUTHOR_RUNNING)}}}",
        rf"\institute{{{institute_tex}}}",
        "",
        r"\maketitle",
        "",
        r"\begin{abstract}",
        latex_escape(ABSTRACT),
        "",
        r"\keywords{" + r" \and ".join(latex_escape(k) for k in KEYWORDS) + "}",
        r"\end{abstract}",
        "",
    ]

    for section in SECTIONS:
        lines.append(rf"\section{{{latex_escape(section['title'])}}}")
        for block in section["blocks"]:
            if isinstance(block, str):
                lines.append(latex_text(block))
                lines.append("")
            elif "subsection" in block:
                lines.append(rf"\subsection{{{latex_escape(block['subsection'])}}}")
            elif "table" in block:
                lines.append(build_latex_table(block["table"]))
                lines.append("")
            elif "figure" in block:
                lines.append(build_latex_figure(block["figure"]))
                lines.append("")
            elif "equation" in block:
                lines.append(r"\begin{equation}")
                lines.append(block["equation"])
                lines.append(r"\end{equation}")
                lines.append("")

    lines.extend(
        [
            r"\section*{Acknowledgements}",
            (
                "The authors thank the annotators, domain reviewers, and technical "
                "contributors who supported the VaccineNLP study, software prototype, "
                "and supporting project materials."
            ),
            "",
            r"\bibliographystyle{splncs04}",
            r"\bibliography{references}",
            "",
            rf"\section*{{{latex_escape(CONTACT_SECTION_TITLE)}}}",
        ]
    )
    for contact_paragraph in CONTACT_PARAGRAPHS:
        contact_text = latex_escape(contact_paragraph)
        for email in CONTACT_EMAILS:
            contact_text = contact_text.replace(email, rf"\texttt{{{email}}}")
        lines.extend([r"\noindent " + contact_text, ""])
    lines.extend([r"\end{document}", ""])
    return "\n".join(lines)


def set_run_font(run, name="Times New Roman", size=10, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def set_cell_margins(cell, top=65, start=60, bottom=65, end=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell, no_wrap=True):
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap_node = tc_pr.first_child_found_in("w:noWrap")
    if no_wrap and no_wrap_node is None:
        no_wrap_node = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap_node)
    elif not no_wrap and no_wrap_node is not None:
        tc_pr.remove(no_wrap_node)


def set_cell_borders(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    def set_edge(edge_name, spec):
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        if spec is None:
            edge.set(qn("w:val"), "nil")
        else:
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), str(spec.get("sz", 6)))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), spec.get("color", "000000"))

    set_edge("top", top)
    set_edge("bottom", bottom)
    set_edge("start", None)
    set_edge("end", None)
    set_edge("insideH", None)
    set_edge("insideV", None)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def set_cell_text(cell, text, bold=False, size=8.0, align="left"):
    paragraph = clear_cell(cell)
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def configure_doc_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(5.2)
    section.bottom_margin = Cm(5.2)
    section.left_margin = Cm(4.4)
    section.right_margin = Cm(4.4)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, before, after in [
        ("Heading 1", 12, 10, 4),
        ("Heading 2", 10, 6, 3),
        ("Heading 3", 10, 4, 2),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def apply_compact_spacing(doc):
    style_settings = {
        "Normal": (0, 1, 1.0),
        "p1a": (0, 1, 1.0),
        "abstract": (0, 3, 1.0),
        "keywords": (0, 4, 1.0),
        "address": (0, 1, 1.0),
        "author": (0, 2, 1.0),
        "papertitle": (0, 5, 1.0),
        "heading1": (8, 3, 1.0),
        "heading2": (5, 2, 1.0),
        "tablecaption": (4, 1, 1.0),
        "figurecaption": (1, 3, 1.0),
        "referenceitem": (0, 1, 1.0),
        "equation": (1, 2, 1.0),
    }
    for style_name, (before, after, line_spacing) in style_settings.items():
        if style_exists(doc, style_name):
            fmt = doc.styles[style_name].paragraph_format
            fmt.space_before = Pt(before)
            fmt.space_after = Pt(after)
            fmt.line_spacing = line_spacing
            if style_name == "Normal":
                fmt.first_line_indent = Cm(0.35)
            else:
                fmt.first_line_indent = Pt(0)


def add_paragraph(
    doc,
    text,
    size=10,
    bold_prefix=None,
    style=None,
    first_line_indent=False,
    use_template_styles=False,
):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if not use_template_styles:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.first_line_indent = Cm(0.35) if first_line_indent else Pt(0)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        if use_template_styles:
            run.bold = True
        else:
            set_run_font(run, size=size, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        if not use_template_styles:
            set_run_font(rest, size=size)
    else:
        run = paragraph.add_run(text)
        if not use_template_styles:
            set_run_font(run, size=size)
    return paragraph


def add_equation(doc, equation, use_template_styles=False):
    paragraph = doc.add_paragraph(style="equation") if "equation" in [s.name for s in doc.styles] else doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not use_template_styles:
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
    text = (
        "L = 1.2 L_misinfo + L_stance + L_sentiment."
        if "misinfo" in equation
        else equation
    )
    run = paragraph.add_run(text)
    if not use_template_styles:
        set_run_font(run, size=10)


def add_caption(doc, number, caption, use_template_styles=False):
    paragraph = (
        doc.add_paragraph(style="tablecaption")
        if "tablecaption" in [s.name for s in doc.styles]
        else doc.add_paragraph()
    )
    if not use_template_styles:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
    run_num = paragraph.add_run(f"Table {number}. ")
    if use_template_styles:
        run_num.bold = True
        run_num.italic = True
    else:
        set_run_font(run_num, size=9, bold=True, italic=True)
    run_caption = paragraph.add_run(caption)
    if use_template_styles:
        run_caption.italic = True
    else:
        set_run_font(run_caption, size=9, italic=True)


def add_note(doc, note):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run_label = paragraph.add_run("Note. ")
    set_run_font(run_label, size=8, italic=True)
    run = paragraph.add_run(note)
    set_run_font(run, size=8)


def add_figure_caption(doc, number, caption, use_template_styles=False):
    paragraph = (
        doc.add_paragraph(style="figurecaption")
        if "figurecaption" in [s.name for s in doc.styles]
        else doc.add_paragraph()
    )
    if not use_template_styles:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
    run_num = paragraph.add_run(f"Fig. {number}. ")
    if use_template_styles:
        run_num.bold = True
    else:
        set_run_font(run_num, size=9, bold=True)
    run_caption = paragraph.add_run(caption)
    if not use_template_styles:
        set_run_font(run_caption, size=9)


def add_figure(doc, figure_key, number, figure_paths, use_template_styles=False):
    figure = FIGURES[figure_key]
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run()
    inline_shape = run.add_picture(
        str(figure_paths[figure_key]), width=Cm(figure["docx_width_cm"])
    )
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("name", f"Fig. {number}")
    doc_pr.set("title", figure["caption"])
    doc_pr.set("descr", figure["alt"])
    add_figure_caption(doc, number, figure["caption"], use_template_styles=use_template_styles)


def style_exists(doc, name):
    return any(style.name == name for style in doc.styles)


def add_styled_text(doc, text, style_name, align=None, use_template_styles=False):
    paragraph = doc.add_paragraph(style=style_name if style_exists(doc, style_name) else None)
    if align is not None:
        paragraph.alignment = align
    if not use_template_styles:
        paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(text)
    return paragraph


def add_author_block(doc, use_template_styles=False):
    paragraph = doc.add_paragraph(style="author" if style_exists(doc, "author") else None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not use_template_styles:
        paragraph.paragraph_format.line_spacing = 1.0
    for idx, author in enumerate(AUTHORS):
        if idx:
            paragraph.add_run(", " if idx < len(AUTHORS) - 1 else ", and ")
        paragraph.add_run(author["name"])
        marker = paragraph.add_run(author["inst"])
        marker.font.superscript = True


def add_institution_block(doc, use_template_styles=False):
    for idx, institution in enumerate(INSTITUTIONS, 1):
        paragraph = doc.add_paragraph(style="address" if style_exists(doc, "address") else None)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not use_template_styles:
            paragraph.paragraph_format.line_spacing = 1.0
        marker = paragraph.add_run(str(idx))
        marker.font.superscript = True
        paragraph.add_run(f" {institution}")


def clear_part_content(part):
    element = part._element
    for child in list(element):
        element.remove(child)


def add_page_field(paragraph, font_size=8.0, bold=False):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run = paragraph.add_run()
    set_run_font(begin_run, size=font_size, bold=bold)
    begin_run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run = paragraph.add_run()
    set_run_font(instr_run, size=font_size, bold=bold)
    instr_run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run = paragraph.add_run()
    set_run_font(separate_run, size=font_size, bold=bold)
    separate_run._r.append(separate)

    result_run = paragraph.add_run("1")
    set_run_font(result_run, size=font_size, bold=bold)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run = paragraph.add_run()
    set_run_font(end_run, size=font_size, bold=bold)
    end_run._r.append(end)


def configure_running_headers(doc):
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.header_distance = Cm(1.25)
        text_width = section.page_width - section.left_margin - section.right_margin

        clear_part_content(section.first_page_header)

        clear_part_content(section.even_page_header)
        paragraph = section.even_page_header.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(0.7), WD_TAB_ALIGNMENT.LEFT)
        add_page_field(paragraph, bold=True)
        paragraph.add_run("\t")
        run = paragraph.add_run(AUTHOR_RUNNING)
        set_run_font(run, size=8.0, bold=True)

        clear_part_content(section.header)
        paragraph = section.header.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
        run = paragraph.add_run(TITLE_RUNNING)
        set_run_font(run, size=8.0, bold=True)
        paragraph.add_run("\t")
        add_page_field(paragraph, bold=True)


def normalize_widths(widths):
    total = sum(widths)
    target = 6900
    scaled = [int(width * target / total) for width in widths]
    scaled[-1] += target - sum(scaled)
    return scaled


def add_table(doc, table_key, number, use_template_styles=False):
    table_def = TABLES[table_key]
    add_caption(doc, number, table_def["caption"], use_template_styles=use_template_styles)
    cols = len(table_def["headers"])
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    widths = normalize_widths(table_def["widths"])
    set_table_width(table, widths)

    font_size = table_def.get("font_size", 7.8)
    aligns = table_def["align"]
    nowrap_cols = set(table_def.get("nowrap_cols", []))
    cell_margins = table_def.get("cell_margins", {})
    header_top = {"sz": 10}
    header_bottom = {"sz": 7}
    bottom_rule = {"sz": 10}

    mark_header_row(table.rows[0])
    for idx, header in enumerate(table_def["headers"]):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell, **cell_margins)
        set_cell_no_wrap(cell, idx in nowrap_cols)
        set_cell_borders(cell, top=header_top, bottom=header_bottom)
        set_cell_text(cell, header, bold=True, size=font_size, align=aligns[idx])

    for row_data in table_def["rows"]:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, **cell_margins)
            set_cell_no_wrap(cell, idx in nowrap_cols)
            set_cell_borders(cell)
            set_cell_text(cell, value, size=font_size, align=aligns[idx])

    for cell in table.rows[-1].cells:
        set_cell_borders(cell, bottom=bottom_rule)

    if table_def.get("note"):
        add_note(doc, table_def["note"])
    else:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)


def sanitize_springer_docm_template():
    if not DOCX_TEMPLATE.exists():
        return None
    cache_dir = ROOT / ".cache"
    cache_dir.mkdir(exist_ok=True)
    template_path = cache_dir / "splnproc1703_sanitized_template.docx"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ET.register_namespace("", rel_ns)
    ET.register_namespace("", ct_ns)
    with ZipFile(DOCX_TEMPLATE, "r") as zin, ZipFile(template_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            name = item.filename
            if name.startswith("customUI/") or name == "word/vbaProject.bin":
                continue
            data = zin.read(name)
            if name == "_rels/.rels":
                root = ET.fromstring(data)
                for rel in list(root):
                    if "customUI" in rel.attrib.get("Target", ""):
                        root.remove(rel)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "word/_rels/document.xml.rels":
                root = ET.fromstring(data.lstrip(b"\xef\xbb\xbf"))
                for rel in list(root):
                    if (
                        "vbaProject" in rel.attrib.get("Type", "")
                        or rel.attrib.get("Target") == "vbaProject.bin"
                    ):
                        root.remove(rel)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "[Content_Types].xml":
                root = ET.fromstring(data)
                for el in list(root):
                    if (
                        el.attrib.get("Extension") == "bin"
                        or "vbaProject" in el.attrib.get("ContentType", "")
                    ):
                        root.remove(el)
                    if el.attrib.get("PartName") == "/word/document.xml":
                        el.set(
                            "ContentType",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                        )
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(item, data)
    return template_path


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_docx(path, figure_paths):
    template_path = sanitize_springer_docm_template()
    doc = Document(str(template_path)) if template_path else Document()
    use_template_styles = template_path is not None
    clear_document_body(doc)
    if not use_template_styles:
        configure_doc_styles(doc)
        apply_compact_spacing(doc)

    add_styled_text(doc, TITLE, "papertitle", WD_ALIGN_PARAGRAPH.CENTER, use_template_styles)
    add_author_block(doc, use_template_styles)
    add_institution_block(doc, use_template_styles)
    add_styled_text(
        doc,
        "Abstract. " + ABSTRACT,
        "abstract",
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        use_template_styles,
    )
    add_styled_text(
        doc,
        "Keywords: " + "; ".join(KEYWORDS) + ".",
        "keywords",
        WD_ALIGN_PARAGRAPH.LEFT,
        use_template_styles,
    )

    table_counter = 0
    figure_counter = 0
    first_body_after_heading = False
    for section in SECTIONS:
        add_styled_text(doc, section["title"], "heading1", use_template_styles=use_template_styles)
        first_body_after_heading = True
        for block in section["blocks"]:
            if isinstance(block, str):
                is_first_body = first_body_after_heading
                style_name = "p1a" if is_first_body else "Normal"
                add_paragraph(
                    doc,
                    block,
                    style=style_name if style_exists(doc, style_name) else None,
                    first_line_indent=not is_first_body,
                    use_template_styles=use_template_styles,
                )
                first_body_after_heading = False
            elif "subsection" in block:
                add_styled_text(
                    doc,
                    block["subsection"],
                    "heading2",
                    use_template_styles=use_template_styles,
                )
                first_body_after_heading = True
            elif "table" in block:
                table_counter += 1
                add_table(doc, block["table"], table_counter, use_template_styles=use_template_styles)
                first_body_after_heading = True
            elif "figure" in block:
                figure_counter += 1
                add_figure(
                    doc,
                    block["figure"],
                    figure_counter,
                    figure_paths,
                    use_template_styles=use_template_styles,
                )
                first_body_after_heading = True
            elif "equation" in block:
                add_equation(doc, block["equation"], use_template_styles=use_template_styles)
                first_body_after_heading = True

    add_styled_text(doc, "Acknowledgements", "heading1", use_template_styles=use_template_styles)
    add_paragraph(
        doc,
        "The authors thank the annotators, domain reviewers, and technical contributors "
        "who supported the VaccineNLP study, software prototype, and supporting "
        "project materials.",
        style="p1a" if style_exists(doc, "p1a") else None,
        first_line_indent=False,
        use_template_styles=use_template_styles,
    )

    add_styled_text(doc, "References", "heading1", use_template_styles=use_template_styles)
    for idx, ref in enumerate(REFERENCES, 1):
        paragraph = (
            doc.add_paragraph(style="referenceitem")
            if style_exists(doc, "referenceitem")
            else doc.add_paragraph()
        )
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(f"[{idx}] {ref}")
        if not use_template_styles:
            set_run_font(run, size=8.5)

    add_styled_text(
        doc,
        CONTACT_SECTION_TITLE,
        "heading1",
        use_template_styles=use_template_styles,
    )
    for idx, paragraph_text in enumerate(CONTACT_PARAGRAPHS):
        add_paragraph(
            doc,
            paragraph_text,
            style=("p1a" if idx == 0 and style_exists(doc, "p1a") else "Normal"),
            first_line_indent=idx != 0,
            use_template_styles=use_template_styles,
        )

    props = doc.core_properties
    props.title = TITLE
    props.author = "VaccineNLP Project Team"
    props.subject = "FISAT full paper"
    props.keywords = ", ".join(KEYWORDS)
    configure_running_headers(doc)
    doc.save(path)


def build_zip(zip_path, paths):
    if zip_path.exists():
        zip_path.unlink()
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
        for path in paths:
            if path.parent.name == "figures":
                zf.write(path, arcname=f"figures/{path.name}")
            else:
                zf.write(path, arcname=path.name)


def generate_architecture_figure(path):
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1800, 1780
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    def font(size, bold=False):
        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                return ImageFont.truetype(candidate, size)
        return ImageFont.load_default()

    title_font = font(34, True)
    body_font = font(28)
    small_font = font(24)
    note_font = font(24)

    def rounded_box(xy, fill, outline, radius=24, width_line=4):
        draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width_line)

    def center_text(xy, text, fnt, fill="#222222", spacing=8):
        x1, y1, x2, y2 = xy
        lines = text.split("\n")
        heights = []
        widths = []
        for line in lines:
            box = draw.textbbox((0, 0), line, font=fnt)
            widths.append(box[2] - box[0])
            heights.append(box[3] - box[1])
        total_h = sum(heights) + spacing * (len(lines) - 1)
        y = y1 + (y2 - y1 - total_h) / 2
        for line, w, h in zip(lines, widths, heights):
            draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
            y += h + spacing

    def arrow(start, end, fill="#444444", width_line=6):
        draw.line([start, end], fill=fill, width=width_line)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) < abs(ey - sy):
            direction = 1 if ey > sy else -1
            pts = [(ex, ey), (ex - 18, ey - 34 * direction), (ex + 18, ey - 34 * direction)]
        else:
            direction = 1 if ex > sx else -1
            pts = [(ex, ey), (ex - 34 * direction, ey - 18), (ex - 34 * direction, ey + 18)]
        draw.polygon(pts, fill=fill)

    rounded_box((570, 55, 1230, 190), "#F5F5F5", "#777777")
    center_text((570, 55, 1230, 190), "Input text\n(Vietnamese vaccine post/comment)", body_font)
    arrow((900, 190), (900, 255))

    rounded_box((205, 255, 1595, 410), "#FCE7C7", "#DD8A2B")
    center_text(
        (205, 255, 1595, 410),
        "Four-layer noise-resistance preprocessing\ncharacters -> words -> semantics -> embedding anomaly",
        body_font,
    )
    arrow((900, 410), (900, 485))

    rounded_box((500, 485, 1300, 685), "#DCECF8", "#2E75B6")
    center_text(
        (500, 485, 1300, 685),
        "TIER 1: PhoBERT-v2 Classification Engine\nclassifies all inputs\nmisinformation - stance - sentiment\n+ confidence",
        title_font,
    )

    draw.text((260, 745), "~80% clear cases", font=small_font, fill="#2E75B6")
    arrow((510, 620), (330, 1190), fill="#2E75B6", width_line=6)
    draw.text((930, 760), "~20% difficult cases\n(sarcasm, ambiguity, long text)", font=small_font, fill="#7F5AC6")
    arrow((900, 685), (900, 830), fill="#7F5AC6", width_line=6)

    rounded_box((500, 830, 1300, 1030), "#E7DDF5", "#7F5AC6")
    center_text(
        (500, 830, 1300, 1030),
        "TIER 2: Gemma-4 4B Explanation Engine\ngenerates reasoning for difficult cases",
        title_font,
    )

    rounded_box((1325, 850, 1760, 1040), "#DCEED6", "#5D9B58")
    center_text(
        (1325, 850, 1760, 1040),
        "Evidence checking (RAG)\nEurope PMC - OpenAlex\nSemantic Scholar - newspapers\nWHO / CDC / Ministry of Health",
        small_font,
    )
    arrow((1325, 945), (1300, 945), fill="#5D9B58", width_line=6)
    arrow((900, 1030), (900, 1190))

    rounded_box((255, 1190, 1545, 1420), "#D9F0ED", "#2C908C")
    center_text(
        (255, 1190, 1545, 1420),
        "COGNITIVE INTEGRITY\nNo automatic true/false judgment without evidence\nconfidence and highlighted text guide human review",
        title_font,
    )
    arrow((900, 1420), (900, 1500))

    rounded_box((500, 1500, 1300, 1690), "#F4F4F4", "#666666")
    center_text(
        (500, 1500, 1300, 1690),
        "Output\nwarning signal + explanation + evidence status\n(not an autonomous medical verdict)",
        body_font,
    )
    image.save(path)


def generate_per_class_f1_figure(path):
    from PIL import Image, ImageDraw, ImageFont

    def font(size, bold=False):
        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                return ImageFont.truetype(candidate, size)
        return ImageFont.load_default()

    models = ["XLM-R-v1", "PhoBERT-v2", "Gemma-4 4B"]
    tasks = [
        ("Misinformation", ["Misinformation", "Correct"], [[0.5079, 0.8997], [0.5075, 0.8918], [0.4444, 0.8309]], [28, 158]),
        ("Stance", ["Supportive", "Opposing", "Neutral"], [[0.5495, 0.6387, 0.6790], [0.5934, 0.6612, 0.7375], [0.4528, 0.6905, 0.7360]], [54, 48, 84]),
        ("Sentiment", ["Negative", "Neutral", "Positive"], [[0.7682, 0.7162, 0.5753], [0.8000, 0.7917, 0.5882], [0.8039, 0.8034, 0.7027]], [71, 75, 40]),
    ]
    colors = ["#5B84C9", "#F08A45", "#82B65D"]
    image = Image.new("RGB", (2684, 894), "white")
    draw = ImageDraw.Draw(image)
    title_f = font(38, True)
    axis_f = font(24)
    small_f = font(20)
    heading_f = font(30, True)
    draw.text((790, 18), "VaccineNLP - Per-Class F1 on the Gold Test Set", font=title_f, fill="black")

    def vertical_label(text, x, y, fnt, fill="#222222"):
        bbox = draw.textbbox((0, 0), text, font=fnt)
        label = Image.new("RGBA", (bbox[2] - bbox[0] + 10, bbox[3] - bbox[1] + 10), (255, 255, 255, 0))
        label_draw = ImageDraw.Draw(label)
        label_draw.text((5 - bbox[0], 5 - bbox[1]), text, font=fnt, fill=fill)
        label = label.rotate(90, expand=True)
        image.paste(label, (int(x), int(y - label.height / 2)), label)

    panel_w = 820
    panel_h = 620
    lefts = [80, 930, 1780]
    top = 110
    for pidx, (title, classes, data, support) in enumerate(tasks):
        left = lefts[pidx]
        chart = (left, top + 70, left + panel_w, top + panel_h)
        draw.text((left + panel_w // 2 - 85, top), title, font=heading_f, fill="black")
        x0, y0, x1, y1 = chart
        draw.line((x0, y1, x1, y1), fill="#222222", width=2)
        draw.line((x0, y0, x0, y1), fill="#222222", width=2)
        for tick in [0, 0.2, 0.4, 0.6, 0.8, 1.0]:
            y = y1 - int((y1 - y0) * tick)
            draw.line((x0, y, x1, y), fill="#E6E6E6", width=1)
            draw.text((x0 - 46, y - 12), f"{tick:.1f}", font=small_f, fill="#333333")
        if pidx == 0:
            vertical_label("F1 score", x0 - 76, (y0 + y1) / 2, axis_f)
        n_classes = len(classes)
        group_w = (x1 - x0) / n_classes
        bar_w = 34 if n_classes == 3 else 46
        for cidx, cls in enumerate(classes):
            cx = x0 + group_w * cidx + group_w / 2
            for midx, color in enumerate(colors):
                value = data[midx][cidx]
                bx = int(cx + (midx - 1) * (bar_w + 8))
                by = y1 - int((y1 - y0) * value)
                draw.rectangle((bx - bar_w // 2, by, bx + bar_w // 2, y1), fill=color)
            label_w = draw.textlength(cls, font=small_f)
            draw.text((cx - label_w / 2, y1 + 12), cls, font=small_f, fill="#222222")
            support_label = f"n={support[cidx]}"
            sw = draw.textlength(support_label, font=small_f)
            draw.text((cx - sw / 2, y1 + 48), support_label, font=small_f, fill="#777777")
        legend_x = x1 - 185
        legend_y = y0 + 12
        for midx, (model, color) in enumerate(zip(models, colors)):
            yy = legend_y + midx * 30
            draw.rectangle((legend_x, yy, legend_x + 24, yy + 14), fill=color)
            draw.text((legend_x + 34, yy - 4), model, font=small_f, fill="#222222")
    image.save(path)


def generate_calibration_figure(path):
    from PIL import Image, ImageDraw, ImageFont

    def font(size, bold=False):
        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                return ImageFont.truetype(candidate, size)
        return ImageFont.load_default()

    panels = [
        ("Misinformation", 0.123, 0.054, 1.82, [0.55, 0.66, 0.75, 0.86, 0.98], [0.25, 0.67, 0.25, 0.67, 0.88], [0.56, 0.66, 0.76, 0.87, 0.94], [0.57, 0.36, 0.72, 0.86, 0.89]),
        ("Stance", 0.198, 0.093, 1.67, [0.39, 0.45, 0.55, 0.66, 0.76, 0.86, 0.97], [1.00, 0.50, 0.33, 0.33, 0.71, 0.50, 0.79], [0.38, 0.47, 0.56, 0.65, 0.77, 0.86, 0.92], [0.50, 0.37, 0.42, 0.67, 0.62, 0.74, 0.91]),
        ("Sentiment", 0.144, 0.081, 1.35, [0.39, 0.47, 0.56, 0.66, 0.77, 0.87, 0.95], [1.00, 0.25, 0.50, 0.37, 0.67, 0.70, 0.84], [0.38, 0.48, 0.54, 0.66, 0.77, 0.87, 0.93], [0.67, 0.40, 0.25, 0.63, 0.68, 0.83, 0.84]),
    ]
    image = Image.new("RGB", (2685, 741), "white")
    draw = ImageDraw.Draw(image)
    title_f = font(36, True)
    heading_f = font(30, True)
    small_f = font(20)
    axis_f = font(24)
    draw.text((710, 18), "PhoBERT-v2 Reliability Diagrams: Confidence vs. Accuracy", font=title_f, fill="black")
    lefts = [75, 955, 1835]
    top = 115
    panel_w, panel_h = 760, 500

    def vertical_label(text, x, y, fnt, fill="#222222"):
        bbox = draw.textbbox((0, 0), text, font=fnt)
        label = Image.new("RGBA", (bbox[2] - bbox[0] + 10, bbox[3] - bbox[1] + 10), (255, 255, 255, 0))
        label_draw = ImageDraw.Draw(label)
        label_draw.text((5 - bbox[0], 5 - bbox[1]), text, font=fnt, fill=fill)
        label = label.rotate(90, expand=True)
        image.paste(label, (int(x), int(y - label.height / 2)), label)

    def map_point(x, y, box):
        x0, y0, x1, y1 = box
        return (x0 + int((x1 - x0) * x), y1 - int((y1 - y0) * y))

    for pidx, (left, (title, raw_ece, cal_ece, temp, raw_x, raw_y, cal_x, cal_y)) in enumerate(zip(lefts, panels)):
        chart = (left, top + 70, left + panel_w, top + panel_h)
        x0, y0, x1, y1 = chart
        draw.text((left + panel_w // 2 - 80, top), title, font=heading_f, fill="black")
        draw.rectangle(chart, outline="#222222", width=2)
        for tick in [0, 0.2, 0.4, 0.6, 0.8, 1.0]:
            x = x0 + int((x1 - x0) * tick)
            y = y1 - int((y1 - y0) * tick)
            draw.line((x, y0, x, y1), fill="#E6E6E6", width=1)
            draw.line((x0, y, x1, y), fill="#E6E6E6", width=1)
            draw.text((x - 10, y1 + 12), f"{tick:.1f}", font=small_f, fill="#333333")
            draw.text((x0 - 48, y - 12), f"{tick:.1f}", font=small_f, fill="#333333")
        draw.line((map_point(0, 0, chart), map_point(1, 1, chart)), fill="#888888", width=3)
        draw.line((map_point(0, 0, chart), map_point(1, 1, chart)), fill="#888888", width=3)
        raw_pts = [map_point(x, y, chart) for x, y in zip(raw_x, raw_y)]
        cal_pts = [map_point(x, y, chart) for x, y in zip(cal_x, cal_y)]
        draw.line(raw_pts, fill="#E64B3C", width=4)
        draw.line(cal_pts, fill="#55DDBB", width=4)
        for px, py in raw_pts:
            draw.ellipse((px - 8, py - 8, px + 8, py + 8), fill="#E64B3C")
        for px, py in cal_pts:
            draw.rectangle((px - 8, py - 8, px + 8, py + 8), fill="#55DDBB")
        draw.text((x0 + 190, y1 + 48), "Confidence", font=axis_f, fill="#222222")
        if pidx == 0:
            vertical_label("Accuracy", x0 - 72, (y0 + y1) / 2, axis_f)
        legend_x, legend_y = x0 + 20, y0 + 16
        draw.line((legend_x, legend_y + 9, legend_x + 34, legend_y + 9), fill="#888888", width=3)
        draw.text((legend_x + 46, legend_y - 3), "Perfect calibration", font=small_f, fill="#222222")
        draw.line((legend_x, legend_y + 39, legend_x + 34, legend_y + 39), fill="#E64B3C", width=4)
        draw.ellipse((legend_x + 13, legend_y + 31, legend_x + 29, legend_y + 47), fill="#E64B3C")
        draw.text((legend_x + 46, legend_y + 27), f"Raw (ECE={raw_ece:.3f})", font=small_f, fill="#222222")
        draw.line((legend_x, legend_y + 69, legend_x + 34, legend_y + 69), fill="#55DDBB", width=4)
        draw.rectangle((legend_x + 13, legend_y + 61, legend_x + 29, legend_y + 77), fill="#55DDBB")
        draw.text((legend_x + 46, legend_y + 57), f"Calibrated T={temp:.2f} (ECE={cal_ece:.3f})", font=small_f, fill="#222222")
    image.save(path)


def prepare_figure_assets():
    figures_dir = OUT_DIR / "figures"
    figures_dir.mkdir(exist_ok=True)
    figure_paths = {}
    for key, figure in FIGURES.items():
        out_path = figures_dir / figure["filename"]
        if figure.get("generated") == "architecture":
            generate_architecture_figure(out_path)
        elif figure.get("generated") == "per_class_f1":
            generate_per_class_f1_figure(out_path)
        elif figure.get("generated") == "calibration":
            generate_calibration_figure(out_path)
        else:
            source_doc = figure["source_doc"]
            source_media = figure["source_media"]
            if not source_doc.exists():
                raise FileNotFoundError(f"Missing source document: {source_doc}")
            with ZipFile(source_doc) as zf:
                if source_media not in zf.namelist():
                    raise FileNotFoundError(f"Missing {source_media} in {source_doc}")
                data = zf.read(source_media)
            out_path.write_bytes(data)
        figure_paths[key] = out_path
    return figure_paths


def main():
    OUT_DIR.mkdir(exist_ok=True)
    tex_path = OUT_DIR / "VaccineNLP_FISAT_TEMPLATE_WITH_FIGURES.tex"
    bib_path = OUT_DIR / "references.bib"
    docx_path = OUT_DIR / "VaccineNLP_FISAT_TEMPLATE_WITH_FIGURES.docx"
    figure_paths = prepare_figure_assets()
    tex_path.write_text(build_latex(), encoding="utf-8")
    bib_path.write_text(build_bib(), encoding="utf-8")
    build_docx(docx_path, figure_paths)

    copied = []
    for name in ["llncs.cls", "splncs04.bst"]:
        src = LATEX_TEMPLATE_DIR / name
        dst = OUT_DIR / name
        if src.exists():
            copyfile(src, dst)
            copied.append(dst)

    zip_path = OUT_DIR / "VaccineNLP_FISAT_TEMPLATE_WITH_FIGURES_package.zip"
    build_zip(zip_path, [tex_path, bib_path, docx_path, *copied, *figure_paths.values()])

    print(f"Wrote {tex_path}")
    print(f"Wrote {bib_path}")
    print(f"Wrote {docx_path}")
    print(f"Wrote {zip_path}")


if __name__ == "__main__":
    main()
